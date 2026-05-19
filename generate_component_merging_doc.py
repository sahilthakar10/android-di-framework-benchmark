#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

HEADER_BG = RGBColor(0x1A, 0x23, 0x7E)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = HEADER_BG

def set_cell_bg(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    bg = shading.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex})
    shading.append(bg)

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = WHITE
        set_cell_bg(cell, '1A237E')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(cell, 'F5F5F5')

def code_block(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)

def para(text):
    doc.add_paragraph(text)

def bold_para(bold_text, normal_text):
    p = doc.add_paragraph()
    p.add_run(bold_text).bold = True
    p.add_run(normal_text)

# ── Title ──
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Component Merging in DI Frameworks')
r.font.size = Pt(24)
r.bold = True
r.font.color.rgb = HEADER_BG

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Hilt vs kotlin-inject-anvil vs Metro')
r.font.size = Pt(12)
r.font.color.rgb = GRAY

doc.add_paragraph()

# ── What is Component Merging ──
heading('What is Component Merging?')
para(
    'Component merging allows DI bindings to be declared in any Gradle module '
    'and automatically discovered and merged into the final DI graph — without '
    'a central file listing them all.'
)

# ── Hilt ──
heading('How Hilt Does It')
para(
    'Each module declares its bindings with @InstallIn. Hilt\'s processor scans all '
    'modules at compile time and merges them into the target component.'
)

code_block(
    '// :feature-auth module\n'
    '@Module\n'
    '@InstallIn(SingletonComponent::class)\n'
    'object AuthModule {\n'
    '    @Provides @Singleton\n'
    '    fun provideAuthManager(): AuthManager = RealAuthManager()\n'
    '}\n'
    '\n'
    '// :feature-payment module (knows nothing about :feature-auth)\n'
    '@Module\n'
    '@InstallIn(SingletonComponent::class)\n'
    'object PaymentModule {\n'
    '    @Provides @Singleton\n'
    '    fun providePaymentService(): PaymentService = RealPaymentService()\n'
    '}\n'
    '\n'
    '// :app module — no manual wiring needed\n'
    '@HiltAndroidApp\n'
    'class MyApp : Application()\n'
    '// Hilt auto-discovers AuthModule + PaymentModule and merges them'
)

para('Adding a new feature module requires zero changes to the app module.')

# ── kotlin-inject-anvil ──
heading('How kotlin-inject-anvil Does It')
para(
    'Same concept, different annotations. @ContributesBinding declares where a binding '
    'belongs. @MergeComponent collects all contributions.'
)

code_block(
    '// :feature-auth module\n'
    '@ContributesBinding(AppScope::class)\n'
    '@Inject @SingleIn(AppScope::class)\n'
    'class RealAuthManager : AuthManager\n'
    '\n'
    '// :feature-payment module\n'
    '@ContributesBinding(AppScope::class)\n'
    '@Inject @SingleIn(AppScope::class)\n'
    'class RealPaymentService : PaymentService\n'
    '\n'
    '// :app module\n'
    '@MergeComponent(AppScope::class)\n'
    '@SingleIn(AppScope::class)\n'
    'abstract class AppComponent\n'
    '// RealAuthManager and RealPaymentService are auto-discovered'
)

# ── Metro ──
heading('How Metro Does It')
para(
    'Metro does NOT support component merging. Every binding must be explicitly '
    'declared in the @DependencyGraph.'
)

code_block(
    '@DependencyGraph(AppScope::class)\n'
    'interface AppGraph {\n'
    '    @Provides fun bindAuth(impl: RealAuthManager): AuthManager = impl\n'
    '    @Provides fun bindPayment(impl: RealPaymentService): PaymentService = impl\n'
    '    @Provides fun bindCart(impl: RealCartRepository): CartRepository = impl\n'
    '    // ... every binding must be listed here\n'
    '}'
)

para('For multi-module setups, Metro uses @Includes to compose graphs:')

code_block(
    '@DependencyGraph(AppScope::class)\n'
    'interface AppGraph {\n'
    '    @DependencyGraph.Factory\n'
    '    fun interface Factory {\n'
    '        fun create(\n'
    '            @Includes authGraph: AuthGraph,\n'
    '            @Includes paymentGraph: PaymentGraph\n'
    '        ): AppGraph\n'
    '    }\n'
    '}'
)

para('Each sub-graph is still explicitly included — no auto-discovery.')

# ── Comparison Table ──
heading('Comparison')

table(
    ['Aspect', 'Hilt', 'kotlin-inject-anvil', 'Metro'],
    [
        ['Auto-discovery', '@InstallIn', '@ContributesTo / @ContributesBinding', 'Not supported'],
        ['Adding a new module', 'Add @InstallIn in new module', 'Add @ContributesBinding in new module', 'Update central @DependencyGraph'],
        ['Central "god file"', 'No', 'No', 'Yes'],
        ['Merge conflicts at scale', 'Rare', 'Rare', 'Common'],
    ]
)

doc.add_paragraph()

# ── When It Matters ──
heading('When It Matters')

bold_para('Small-to-medium apps (< 15 modules): ',
    'Component merging is nice-to-have but not critical. Metro\'s explicit graph '
    'is readable and manageable.')

bold_para('Large apps (50+ modules, multiple teams): ',
    'Component merging becomes important. Without it, the central graph file becomes '
    'a bottleneck — every team touches it, merge conflicts increase, and team autonomy decreases.')

# ── Why Metro Doesn't Support It ──
heading('Why Metro Doesn\'t Support It')

para(
    'Metro is a compiler plugin (FIR + IR). Component merging requires cross-module '
    'discovery at compile time — scanning annotations across all dependencies. '
    'KSP processors (Hilt, kotlin-inject-anvil) do this naturally because they run '
    'as annotation processors with access to all symbols. A compiler plugin operates '
    'within a single module\'s compilation, making cross-module discovery architecturally harder.'
)

para('Metro\'s trade-off: faster compilation and zero generated files, at the cost of explicit wiring.')

# ── Save ──
out = '/Users/sahilthakar/AndroidStudioProjects/BenchMarking/docs/Component_Merging_in_DI_Frameworks.docx'
doc.save(out)
print(f'Saved: {out}')
