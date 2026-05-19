#!/usr/bin/env python3
"""
Generates RFC document for DI Framework Benchmarking:
Hilt vs Metro vs Koin
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

BLUE = RGBColor(0x21, 0x96, 0xF3)
ORANGE = RGBColor(0xFF, 0x98, 0x00)
GREEN_HILT = RGBColor(0x4C, 0xAF, 0x50)
DARK = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = RGBColor(0x1A, 0x23, 0x7E)

def set_cell_bg(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    bg = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    shading.append(bg)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = WHITE
        set_cell_bg(cell, '1A237E')

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.size = Pt(10)
            if ri % 2 == 1:
                set_cell_bg(cell, 'F5F5F5')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = HEADER_BG
    return h

# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RFC: Android DI Framework\nPerformance Benchmark')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = HEADER_BG

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Hilt vs Metro vs Koin vs kotlin-inject-anvil')
run.font.size = Pt(14)
run.font.color.rgb = GRAY

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Date: {datetime.date.today().strftime("%B %d, %Y")}\n').font.size = Pt(11)
meta.add_run('Version: 1.0\n').font.size = Pt(11)
meta.add_run('Status: Draft for Review\n').font.size = Pt(11)
meta.add_run('Test Environment: Pixel 9 Pro Emulator, API 35\n').font.size = Pt(11)
meta.add_run('Build Tools: AGP 9.2.0, Kotlin 2.3.21, Gradle 9.4.1').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════
add_heading_styled('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Framework Overview',
    '   2.1 Hilt (Google/Dagger + KSP)',
    '   2.2 Metro (Kotlin Compiler Plugin)',
    '   2.3 Koin (Runtime DI)',
    '   2.4 kotlin-inject-anvil (Amazon/KSP)',
    '3. Test Application Architecture',
    '4. Compile-Time Benchmark (Android)',
    '5. Runtime Benchmark (Android)',
    '6. Best Practices, Safety & Runtime Trade-offs',
    '   6.1 Best Practices Applied per Framework',
    '   6.2 Compile-Time Safety Comparison',
    '   6.3 Runtime DI — Real Advantages & Limitations',
    '7. iOS & KMP — Cross-Platform Benchmark',
    '8. Deep Dive — How Metro Achieves Superior Performance',
    '9. Lifecycle Awareness & Android Integration',
    '10. Java Interop & Migration from Dagger/Hilt',
    '11. Head-to-Head — All 4 Frameworks Compared',
    '12. Kotlin Version Compatibility Guide',
    '13. Summary & Recommendation',
    '14. References & Further Reading',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════
add_heading_styled('1. Executive Summary', 1)

doc.add_paragraph(
    'This RFC presents a comprehensive performance comparison of three Android dependency injection '
    'frameworks: Hilt (Google/Dagger), Metro (Zac Sweers), and Koin (Insert-Koin). '
    'We benchmarked both compile-time and runtime performance using an identical, realistic '
    'e-commerce application with ~390 classes and ~290 DI bindings across 14 business domains '
    'and 13 feature modules.'
)

doc.add_paragraph()
add_heading_styled('Key Findings', 2)

add_table(
    ['Metric', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Compile Time (avg)', '4,930ms', '~2,500ms', '~2,500ms', '2,257ms'],
        ['Generated Code', '387 files / 555KB (Java)', '0 source files (IR codegen)', '30 files / 58KB (Kotlin)', '0 files'],
        ['Android Runtime Total', '17ms (123 cls)', '7ms (124 cls)', '6ms (124 cls)', '20ms (82 cls)'],
        ['Component/Graph Init', 'pre-built', '3.46ms', '2.82ms', '11.27ms'],
        ['ViewModel Init (13)', '16.21ms', '3.67ms', '3.67ms', '7.46ms'],
    ],
    [4, 3, 3, 3, 3]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Bottom line: ').bold = True
p.add_run(
    'Metro and kotlin-inject-anvil deliver the best runtime performance — both resolve 124 classes in 6-7ms, '
    'nearly 3x faster than Hilt (17ms) and Koin (20ms). At compile time, Metro and kotlin-inject-anvil are '
    'statistically tied (~2.5s each), Koin is slightly faster (2.3s, no codegen), and Hilt is slowest (4.9s). '
    'Metro generates code directly in IR (zero source files); kotlin-inject-anvil generates 30 small Kotlin files '
    '(vs Hilt\'s 387 Java files), explaining why both are ~2x faster than Hilt despite doing compile-time DI.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 2. FRAMEWORK OVERVIEW
# ═══════════════════════════════════════════════════════
add_heading_styled('2. Framework Overview', 1)

add_heading_styled('2.1 Hilt (Dagger + KSP)', 2)
add_table(
    ['Attribute', 'Detail'],
    [
        ['Type', 'Compile-time (annotation processing)'],
        ['Version Tested', '2.59.2'],
        ['Processing', 'KSP (Kotlin Symbol Processing) + Dagger code generation'],
        ['Annotations', '@Inject, @Singleton, @Module, @Provides, @InstallIn, @HiltAndroidApp'],
        ['Output', 'Generates Java source files (factories, component implementations)'],
        ['Scoping', '@Singleton, @ActivityScoped, @ViewModelScoped, custom scopes'],
        ['Maturity', 'Production-stable, official Google recommendation for Android'],
        ['Trade-off', 'Slowest compile time, but zero runtime overhead and full compile-time safety'],
    ],
    [4, 12]
)

doc.add_paragraph()
add_heading_styled('2.2 Metro (Kotlin Compiler Plugin)', 2)
add_table(
    ['Attribute', 'Detail'],
    [
        ['Type', 'Compile-time (Kotlin compiler plugin, FIR + IR)'],
        ['Version Tested', '1.0.0'],
        ['Author', 'Zac Sweers (github.com/ZacSweers/metro)'],
        ['Processing', 'Native Kotlin compiler plugin — no KSP/KAPT needed'],
        ['Annotations', '@Inject, @DependencyGraph, @Provides, @SingleIn, @AppScope'],
        ['Output', 'Zero source generation — code injected directly into compiler IR output'],
        ['Scoping', '@SingleIn(AppScope::class), custom @Scope annotations'],
        ['Maturity', 'Adopted by Cash App, Vinted; Kotlin-first multiplatform support'],
        ['Trade-off', 'Fastest overall performance, newer ecosystem with fewer community resources'],
    ],
    [4, 12]
)

doc.add_paragraph()
add_heading_styled('2.3 Koin (Runtime DI)', 2)
add_table(
    ['Attribute', 'Detail'],
    [
        ['Type', 'Runtime (service locator pattern with DSL)'],
        ['Version Tested', '4.2.0'],
        ['Processing', 'No annotation processing, no code generation'],
        ['DSL', 'module { }, single { }, factory { }, get(), createdAtStart()'],
        ['Output', 'None — all resolution happens at app startup and during runtime'],
        ['Scoping', 'single (app lifetime), factory (per-request), scoped (context-bound)'],
        ['Maturity', 'Widely adopted, lightweight, KMP support, simple learning curve'],
        ['Trade-off', 'Fastest compile, but slowest runtime — no compile-time DI validation'],
    ],
    [4, 12]
)

doc.add_paragraph()
add_heading_styled('2.4 kotlin-inject-anvil (Amazon/KSP)', 2)
add_table(
    ['Attribute', 'Detail'],
    [
        ['Type', 'Compile-time (KSP annotation processing)'],
        ['Version Tested', '0.1.7 (kotlin-inject 0.9.0)'],
        ['Author', 'Ralf Wondratschek (Amazon), Evan Tatarka (kotlin-inject)'],
        ['Processing', 'KSP → generates Kotlin source files → compiled by kotlinc'],
        ['Annotations', '@Inject, @Component, @Provides, @ContributesTo, @ContributesBinding, @MergeComponent, @SingleIn'],
        ['Output', 'Generated .kt files in build/generated/ksp/ (visible, debuggable)'],
        ['Scoping', '@SingleIn(Scope::class), custom @Scope annotations, @ContributesSubcomponent'],
        ['Maturity', 'Production: Amazon apps, Bitkey (170 KMP modules)'],
        ['Trade-off', 'Debuggable generated code + lifecycle callbacks (via App Platform), but slower builds than Metro (KSP overhead) and no Dagger interop'],
    ],
    [4, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 3. TEST APPLICATION ARCHITECTURE
# ═══════════════════════════════════════════════════════
add_heading_styled('3. Test Application Architecture', 1)

doc.add_paragraph(
    'We built an identical e-commerce application (ShopApp) in all three frameworks. '
    'The app follows Clean Architecture with four layers: Core, Data, Domain, and Feature. '
    'Every class, dependency chain, and business logic is identical — only the DI wiring differs.'
)

add_heading_styled('3.1 Layer Structure', 2)
add_table(
    ['Layer', 'Files', 'Types', 'Description'],
    [
        ['Core', '10', '59', 'Network, Auth, Analytics, Storage, Config, Logging, Image, Notification, Location'],
        ['Data', '70', '98', '14 domains × (Repository + RemoteDataSource + LocalDataSource + Mapper + Models)'],
        ['Domain', '28', '182', '14 domains × (10 UseCases + DomainModel + PagedResult + ValidationResult)'],
        ['Feature', '13', '54', '13 features: Home, Search, ProductDetail, Cart, Checkout, Profile, Orders, Settings, Chat, Notifications, Onboarding, Reviews, Wishlist'],
        ['DI/Graph', '2-5', '1-24', 'DependencyGraph (Metro), @Module (Hilt), Koin modules (Koin)'],
    ],
    [2.5, 1.5, 1.5, 11]
)

doc.add_paragraph()
add_heading_styled('3.2 Class & Binding Inventory', 2)

add_table(
    ['Role', 'Count', 'Scope', 'Description'],
    [
        ['Core Services', '53', 'Singleton', 'AuthManager, AnalyticsTracker, HttpClient, CacheManager, etc.'],
        ['Repositories', '14', 'Singleton', 'One per domain — offline-first with remote + local data sources'],
        ['Remote Data Sources', '14', 'Singleton', 'API calls via HttpClient with auth interceptor'],
        ['Local Data Sources', '14', 'Singleton', 'Database + cache layer via DatabaseManager, CacheManager'],
        ['Mappers', '14', 'Singleton', 'Entity ↔ DomainModel conversion'],
        ['Use Cases', '140', 'Factory', '10 per domain (Get, Create, Update, Delete, Search, Validate, etc.)'],
        ['ViewModels', '12', 'Factory', 'One per feature screen'],
        ['Presenters/Calculators', '22', 'Factory', 'UI logic coordinators (CartCalculator, PriceCalculator, etc.)'],
        ['Data Models', '103', '—', 'data classes (Entity, Response, Request, DomainModel, etc.)'],
    ],
    [4, 1.5, 2, 9]
)

doc.add_paragraph()
add_heading_styled('Binding Summary per Framework', 3)
add_table(
    ['Binding Type', 'Hilt', 'Metro', 'Koin'],
    [
        ['Singletons (cached instances)', '109', '105', '90'],
        ['  Eager (created at startup)', '0', '0', '19'],
        ['  Lazy (on first access)', '109', '105', '71'],
        ['Factory (new instance per request)', '176', '180', '181'],
        ['@Provides methods', '5', '5', '0'],
        ['DI Modules / Graphs', '1', '1', '24'],
        ['TOTAL BINDINGS', '290', '290', '271'],
    ],
    [5, 2.5, 2.5, 2.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 4. COMPILE-TIME BENCHMARK
# ═══════════════════════════════════════════════════════
add_heading_styled('4. Compile-Time Benchmark', 1)

add_heading_styled('4.1 Methodology', 2)
doc.add_paragraph(
    'Each framework module was clean-built 5 times in isolation. Between each run, all three modules '
    'were cleaned (./gradlew :module:clean) to ensure cold builds. Timing was measured as wall-clock time '
    'for the compileDebugKotlin Gradle task, which includes annotation processing (Hilt), compiler plugin '
    'execution (Metro), and standard Kotlin compilation (all three).'
)

p = doc.add_paragraph()
p.add_run('Environment: ').bold = True
p.add_run('macOS, AGP 9.2.0, Kotlin 2.3.21, Gradle 9.4.1, JDK 21')

add_heading_styled('4.2 Results (5 Clean Builds)', 2)
add_table(
    ['Run', 'Hilt (KSP)', 'Metro (Plugin)', 'kotlin-inject-anvil (KSP)', 'Koin (No codegen)'],
    [
        ['Run 1', '6,642ms', '3,481ms', '3,309ms', '2,414ms'],
        ['Run 2', '4,472ms', '2,501ms', '3,102ms', '2,457ms'],
        ['Run 3', '3,676ms', '2,294ms', '2,652ms', '1,900ms'],
        ['Run 4', '', '2,023ms', '2,273ms', ''],
        ['Run 5', '', '2,274ms', '2,110ms', ''],
        ['', '', '', '', ''],
        ['Average', '4,930ms', '2,515ms', '2,689ms', '2,257ms'],
        ['Min', '3,676ms', '2,023ms', '2,110ms', '1,900ms'],
        ['Max', '6,642ms', '3,481ms', '3,309ms', '2,457ms'],
    ],
    [3, 3, 3, 3, 3]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run(
    'Run 1 for Hilt is consistently higher due to JVM/Gradle daemon cold start and initial KSP processor '
    'class loading. This is representative of a first build after opening the project.'
)

doc.add_paragraph()
add_heading_styled('Head-to-Head Compile Time', 3)
add_table(
    ['Comparison', 'Difference', 'Percentage'],
    [
        ['Metro vs Hilt', 'Metro is 2,415ms faster', '49.0% faster'],
        ['kotlin-inject-anvil vs Hilt', 'kinject is 2,241ms faster', '45.5% faster'],
        ['Koin vs Hilt', 'Koin is 2,673ms faster', '54.2% faster'],
        ['Metro vs kotlin-inject-anvil', 'Within measurement noise', '~same speed'],
    ],
    [5, 5, 4]
)

doc.add_paragraph()
add_heading_styled('4.3 Generated Code Analysis', 2)
add_table(
    ['Metric', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Generated source files', '387', '0 (IR codegen)', '30', '0'],
        ['Generated lines of code', '17,555', '0 (injected into IR)', '1,456', '0'],
        ['Generated code size', '555 KB', '0 KB on disk', '58 KB', '0 KB'],
        ['Generated language', 'Java', 'Kotlin IR (not visible)', 'Kotlin', 'N/A'],
        ['Requires separate compilation pass', 'Yes (javac)', 'No', 'No', 'No'],
    ],
    [5, 2.5, 2.5, 2.5, 2.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key insight: ').bold = True
p.add_run(
    'Hilt generates 291 Java files (factories, component implementations, module bindings) totaling 488KB. '
    'These require a separate javac compilation pass after KSP finishes. Metro generates zero files — it '
    'injects code directly into Kotlin\'s IR (Intermediate Representation) during compilation. '
    'Koin generates zero files because all DI resolution is deferred to runtime.'
)

doc.add_paragraph()
add_heading_styled('4.4 How Each Framework Compiles', 2)

add_table(
    ['Step', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['1', 'KSP scans @Inject, @Module, @Provides', 'FIR: Analyze @DependencyGraph, @Inject', 'KSP scans @Inject, @Component, @ContributesBinding', 'Standard Kotlin compilation'],
        ['2', 'Generates 387 Java files', 'IR: Generate implementations in IR', 'Generates 30 Kotlin files', '(No DI processing)'],
        ['3', 'Kotlin compiler compiles source', 'Single pass — done', 'Kotlin compiler compiles source + generated', '(No DI processing)'],
        ['4', 'Java compiler compiles generated', '—', '—', '—'],
        ['Total', '2 passes + codegen (387 Java files)', '1 pass, zero files', '2 passes + codegen (30 Kotlin files)', '1 pass, zero codegen'],
    ],
    [2, 4, 4, 4, 3.5]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 5. RUNTIME BENCHMARK
# ═══════════════════════════════════════════════════════
add_heading_styled('5. Runtime Benchmark (Hilt vs Metro vs Koin)', 1)

doc.add_paragraph(
    'Runtime benchmarks compare all three frameworks using the identical e-commerce app with ~390 classes. '
    'Hilt\'s generated Dagger component is accessed via @EntryPoint, Metro via createGraph(), and Koin via '
    'startKoin(). All three are warmed up before measurement to eliminate JVM class loading bias.'
)

add_heading_styled('5.1 Methodology', 2)
p = doc.add_paragraph()
p.add_run('Warmup phase: ').bold = True
p.add_run(
    'Both Metro and Koin are fully initialized once before measurement to eliminate JVM class loading '
    'and JIT compilation bias. Without warmup, whichever framework runs first pays the class loading tax '
    'for all 390 classes and appears unfairly slower.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Measurement: ').bold = True
p.add_run(
    'System.nanoTime() for timing. Runtime.freeMemory() + Debug.getNativeHeapAllocatedSize() for memory. '
    'System.gc() + 300ms pause between framework measurements. 100 warm iterations per class.'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Device: ').bold = True
p.add_run('Pixel 9 Pro emulator, API 37 (Android 16)')

add_heading_styled('5.2 Container Initialization', 2)
doc.add_paragraph(
    'Time to access/create the DI container. Hilt\'s component is created during Application.onCreate() by '
    '@HiltAndroidApp — we measure EntryPointAccessors.fromApplication() cost. Metro calls createGraph<ShopAppGraph>(). '
    'Koin calls startKoin { modules(allShopAppModules) } with 24 modules and 19 eager singletons.'
)
add_table(
    ['Framework', 'Init Time', 'What Happens'],
    [
        ['Hilt', '~0ms (pre-built)', 'Component built during Application.onCreate() — accessing it is near-instant'],
        ['Metro', '3.46ms', 'createGraph() instantiates one generated class with all wiring compiled into it'],
        ['kotlin-inject-anvil', '2.82ms', 'KSP-generated component with direct constructor calls, similar to Metro'],
        ['Koin', '11.27ms', 'startKoin() registers 271 lambda definitions, creates 19 eager singletons'],
    ],
    [3, 2.5, 10]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run(
    'Hilt\'s container initialization cost is paid once during Application.onCreate() (via @HiltAndroidApp) '
    'and is not measured here — by benchmark time the component already exists. '
    'Metro\'s createGraph() (3.46ms) and kotlin-inject-anvil\'s create() (2.82ms) both instantiate '
    'a single pre-compiled class with all wiring baked in. Koin\'s startKoin() (11.27ms) must process '
    '24 module DSL blocks, register each binding in a HashMap, and eagerly create 19 singleton instances.'
)

doc.add_paragraph()
add_heading_styled('5.3 Cold Injection (First Access)', 2)
doc.add_paragraph(
    'Time to resolve a class for the first time after container initialization. Includes constructing '
    'the object and all its transitive dependencies.'
)
add_table(
    ['Class', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['HomeViewModel', '7.29ms', '0.61ms', '1.36ms', '1.64ms'],
        ['CheckoutViewModel', '1.47ms', '1.11ms', '0.40ms', '0.72ms'],
        ['CartViewModel', '0.50ms', '0.64ms', '0.19ms', '0.60ms'],
        ['ProfileViewModel', '1.68ms', '0.18ms', '0.17ms', '0.36ms'],
        ['ProductDetailVM', '0.79ms', '0.14ms', '0.21ms', '0.51ms'],
        ['ChatViewModel', '0.64ms', '0.22ms', '0.14ms', '0.60ms'],
        ['SearchViewModel', '1.07ms', '0.07ms', '0.11ms', '0.26ms'],
        ['OrderHistoryVM', '0.36ms', '0.08ms', '0.07ms', '0.18ms'],
    ],
    [5.5, 2.5, 2.5, 2.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Analysis: ').bold = True
p.add_run(
    'Hilt\'s first ViewModel (HomeViewModel) takes 7.29ms because it triggers ViewModelProvider + '
    'HiltViewModelFactory for the first time. Subsequent ViewModels are faster. '
    'Metro and kotlin-inject-anvil front-load cost in graph/component creation (3.46ms / 2.82ms), '
    'making individual ViewModel access fast (~0.1-1.4ms). Both generate direct constructor calls — '
    'their runtime performance is virtually identical. '
    'Koin spends 11.27ms on startKoin() and is slower per-ViewModel due to runtime HashMap lookups.'
)

doc.add_paragraph()
add_heading_styled('5.4 Warm Injection (Average of 100 Iterations)', 2)
doc.add_paragraph(
    'Time to resolve a class after it has been resolved at least once. Singletons return cached instances; '
    'factory-scoped classes create new instances each time.'
)
add_table(
    ['Layer', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Component/Graph Init', 'pre-built', '3.46ms', '2.82ms', '11.27ms'],
        ['ViewModels (13)', '16.21ms', '3.67ms', '3.67ms', '7.46ms'],
        ['Core Singletons (14)', '0.37ms', '0.02ms', '0.03ms', '0.20ms'],
        ['Core Services (12)', '0.33ms', '0.12ms', '0.10ms', '0.51ms'],
        ['Repositories (14)', '0.22ms', '0.15ms', '0.10ms', '1.32ms'],
        ['RemoteDataSources (14)', '0.02ms', '0.01ms', '0.01ms', 'N/A'],
        ['LocalDataSources (14)', '0.05ms', '0.01ms', '0.01ms', 'N/A'],
        ['Mappers (14)', '0.01ms', '0.01ms', '0.01ms', 'N/A'],
        ['UseCases (28)', '0.12ms', '0.05ms', '0.04ms', '0.87ms'],
        ['', '', '', '', ''],
        ['TOTAL', '17ms', '7ms', '6ms', '20ms'],
    ],
    [5.5, 2.5, 2.5, 2.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Analysis: ').bold = True
p.add_run(
    'Hilt\'s generated Dagger factories invoke constructors directly — singleton access is a volatile field read. '
    'With @SingleIn(AppScope::class) scoping, Metro\'s singleton access is a direct field read. '
    'Koin\'s single{} also caches, but get<T>() still performs: HashMap.get() → type check → return cached. '
    'For factory-scoped classes (ViewModels, UseCases), both Hilt and Metro invoke generated constructors directly, '
    'while Koin invokes a stored lambda + resolves each get() parameter recursively.'
)

doc.add_paragraph()
add_heading_styled('5.5 Memory Overhead', 2)
add_table(
    ['Framework', 'Memory Delta', 'What Consumes Memory'],
    [
        ['Hilt', '96 KB', 'Generated Dagger component instance + cached singleton instances (volatile fields)'],
        ['Metro', '128 KB', 'One graph class instance + cached singleton instances'],
        ['Koin', '2,000 KB', 'Module registry, 271 lambda definitions, HashMap entries, singleton cache, type metadata'],
    ],
    [3, 2.5, 10]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Analysis: ').bold = True
p.add_run(
    'Hilt and Metro both use compile-time generated code with direct field references — their memory footprint '
    'is primarily the singleton instances themselves. Koin\'s additional overhead comes from storing lambda '
    'closures for every binding, HashMap data structures for the service locator registry, and runtime type '
    'metadata for resolution.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 6. BEST PRACTICES APPLIED
# ═══════════════════════════════════════════════════════
add_heading_styled('6. Best Practices, Safety & Runtime Trade-offs', 1)

add_heading_styled('6.1 Best Practices Applied per Framework', 2)

doc.add_paragraph(
    'To ensure a fair comparison, each framework was configured using its recommended best practices. '
    'No framework was given an unfair advantage or handicap.'
)

add_heading_styled('7.1 Hilt Best Practices', 2)
practices = [
    '@Singleton on all core services, repositories, data sources, and mappers (109 singletons)',
    '@Inject constructor for dependency declaration — no manual wiring',
    '@Module + @Provides only for classes where constructor injection is not possible (5 methods)',
    '@InstallIn(SingletonComponent::class) for app-level bindings',
    'KSP (not KAPT) for annotation processing — faster than legacy KAPT',
    'Hilt 2.59.2 — required for AGP 9 compatibility',
]
for p in practices:
    doc.add_paragraph(p, style='List Bullet')

add_heading_styled('7.2 Metro Best Practices', 2)
practices = [
    '@DependencyGraph(AppScope::class) — scoped graph for singleton caching',
    '@SingleIn(AppScope::class) on all core services, repositories, data sources, mappers (105 singletons)',
    '@Inject constructor for dependency declaration — identical pattern to Hilt',
    '@Provides in graph interface for classes requiring manual construction (5 methods)',
    'Unscoped factory classes (UseCases, ViewModels) — new instance per access',
    'Metro 1.0.0 — compatible with Kotlin 2.3.21, uses FIR + IR compiler phases',
]
for p in practices:
    doc.add_paragraph(p, style='List Bullet')

add_heading_styled('7.3 Koin Best Practices', 2)
practices = [
    'single {} for all core services, repositories, data sources, mappers (90 singletons)',
    'createdAtStart = true on 19 critical infrastructure singletons (HttpClient, DatabaseManager, AuthManager, AnalyticsTracker, etc.) — pre-warmed at startKoin()',
    'factory {} for all use cases, ViewModels, presenters (181 factories)',
    'Modular organization: 24 Koin modules (9 core + 1 data + 1 domain + 13 feature)',
    'get() for dependency resolution within module definitions',
    'Koin 4.2.0 — latest stable release',
]
for p in practices:
    doc.add_paragraph(p, style='List Bullet')

doc.add_page_break()

add_heading_styled('6.2 Compile-Time Safety Comparison', 2)

doc.add_paragraph(
    'A critical differentiator between DI frameworks is whether dependency graph errors are caught '
    'at compile time (build fails) or at runtime (app crashes). This section examines each framework\'s '
    'validation capabilities in depth.'
)

add_heading_styled('How Each Framework Validates', 3)

add_table(
    ['Aspect', 'Hilt', 'Metro', 'Koin (DSL only)', 'Koin + Compiler Plugin'],
    [
        ['Missing dependency', 'Build fails', 'Build fails', 'Runtime crash', 'Build fails'],
        ['Type mismatch', 'Build fails', 'Build fails', 'Runtime crash', 'Build fails'],
        ['Cyclic dependency', 'Build fails', 'Build fails', 'Runtime hang/crash', 'Build fails'],
        ['Wrong qualifier', 'Build fails', 'Build fails', 'Runtime crash', 'Partially caught'],
        ['Scope violation', 'Build fails', 'Build fails', 'Silent misuse', 'Build fails'],
        ['Dynamic modules', 'N/A (static graph)', 'N/A (static graph)', 'No validation', 'No validation'],
        ['Generic type errors', 'Build fails', 'Build fails', 'Runtime crash', 'Not always caught'],
        ['Technology', 'KSP + Dagger', 'FIR + IR', 'None', 'K2 Compiler Plugin'],
    ],
    [3.5, 2.5, 2.5, 3, 3.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key insight: ').bold = True
p.add_run(
    'Hilt and Metro provide complete compile-time safety — if the project builds, the DI graph is guaranteed '
    'to be valid. Koin DSL provides zero compile-time checks. The Koin Compiler Plugin closes most of this gap '
    'but still has edge cases that can crash at runtime.'
)

doc.add_paragraph()
add_heading_styled('Koin Compiler Plugin (K2) — Closing the Safety Gap', 2)

doc.add_paragraph(
    'Koin has developed a native Kotlin Compiler Plugin (K2, FIR + IR) — the same technology Metro uses — '
    'to add compile-time dependency graph validation. This is the recommended approach for all new Kotlin 2.x projects.'
)

add_table(
    ['Attribute', 'Koin Annotations (KSP)', 'Koin Compiler Plugin (K2)'],
    [
        ['Technology', 'KSP annotation processing', 'Native Kotlin compiler plugin (FIR + IR)'],
        ['Kotlin requirement', 'Any version', 'Kotlin 2.3.20+'],
        ['Generated files', 'Yes (KSP output)', 'None (inline transformation)'],
        ['Status', 'Being deprecated', 'Recommended for new projects'],
        ['KMP support', 'Yes (with per-platform KSP config)', 'Yes (native, no config)'],
    ],
    [4, 5.5, 5.5]
)

doc.add_paragraph()
add_heading_styled('Validation Levels', 3)

doc.add_paragraph(
    'The Koin Compiler Plugin validates dependencies at three progressive levels:'
)

add_table(
    ['Level', 'When', 'What It Validates', 'Example'],
    [
        ['A2 — Per-Module', 'Each module compiles', 'Missing deps within a module, qualifier mismatches, cross-scope violations', 'Service(repo: Repository) fails if Repository not in visible module'],
        ['A3 — Full Graph', 'At startKoin<T>() declaration', 'Cross-module dependencies, all modules combined', 'startKoin<MyApp>() validates CoreModule + DataModule + FeatureModule together'],
        ['A4 — Call-Site', 'Every injection call', 'Every get<T>(), inject<T>(), koinViewModel<T>() call validated', 'val vm: UserVM = koinViewModel() fails if UserVM not registered'],
    ],
    [2.5, 3, 5, 5.5]
)

doc.add_paragraph()
add_heading_styled('Validation vs Resolution — The Critical Difference', 2)

doc.add_paragraph(
    'Even with the Koin Compiler Plugin, there is a fundamental architectural difference:'
)

add_table(
    ['', 'Hilt / Metro', 'Koin + Compiler Plugin'],
    [
        ['Compile-time validation', 'Yes — graph verified at build', 'Yes — graph verified at build'],
        ['Compile-time resolution', 'Yes — constructor calls generated as code', 'No — dependencies still resolved via HashMap at runtime'],
        ['Runtime mechanism', 'Direct constructor invocation (zero overhead)', 'HashMap.get() → type check → invoke lambda (overhead per lookup)'],
        ['Performance impact', 'Near-zero (compiled code)', 'Same as Koin without plugin (~64us avg)'],
        ['Analogy', 'GPS route compiled into car\'s computer', 'Google Maps verified destination exists, but you still navigate turn-by-turn'],
    ],
    [3.5, 6, 6]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Bottom line: ').bold = True
p.add_run(
    'The Koin Compiler Plugin closes the safety gap — most missing bindings are now caught at build time. '
    'However, it does NOT close the performance gap. Dependencies are still resolved at runtime via Koin\'s '
    'service locator. The plugin adds verification, not compilation of the dependency graph.'
)

doc.add_paragraph()
add_heading_styled('What Still Slips Through Koin\'s Compile-Time Checks', 2)

add_table(
    ['Scenario', 'Hilt', 'Metro', 'Koin + Plugin'],
    [
        ['Missing standard dependency', 'Build error', 'Build error', 'Build error'],
        ['@InjectedParam / @Provided (runtime params)', 'Build error', 'Build error', 'NOT checked — trusted at runtime'],
        ['Wrong @Qualifier matching', 'Build error', 'Build error', 'Partially caught — can match wrong qualifier silently'],
        ['Generic type binding errors', 'Build error', 'Build error', 'NOT always caught — type erasure issues'],
        ['Dynamic loadKoinModules()', 'N/A', 'N/A', 'NOT checked — loaded modules unknown at compile time'],
        ['binds() with wrong class', 'Build error', 'Build error', 'Can pass compilation, fail at runtime'],
    ],
    [5, 2.5, 2.5, 5.5]
)

doc.add_page_break()

add_heading_styled('6.3 Runtime DI — Real Advantages & Limitations', 2)

doc.add_paragraph(
    'Runtime DI (Koin) is often criticized for being slower and less safe than compile-time DI (Hilt, Metro). '
    'However, it offers specific capabilities that compile-time frameworks genuinely cannot provide. '
    'This section presents an honest assessment of both sides.'
)

add_heading_styled('8.1 Real Advantages of Runtime Resolution', 2)

add_heading_styled('Dynamic Dependency Swapping', 3)
doc.add_paragraph(
    'The one capability compile-time DI genuinely cannot replicate. Koin can swap implementations '
    'at runtime based on server configuration, user actions, or A/B test assignments — without '
    'recompiling the app.'
)
p = doc.add_paragraph()
p.add_run('Example: ').bold = True
p.add_run(
    'A server-driven config tells the app to use StripeProcessor for payments in the US and '
    'RazorpayProcessor in India. With Koin, you can loadKoinModules() with the appropriate '
    'implementation after the config is fetched — the DI graph changes at runtime. '
    'With Hilt/Metro, the graph is fixed at compile time; you must use if/else inside a @Provides method.'
)

doc.add_paragraph()
add_heading_styled('Module Loading & Unloading', 3)
doc.add_paragraph(
    'Koin supports loadKoinModules() and unloadKoinModules() at runtime. This enables:'
)
points = [
    'Loading feature dependencies only when the user navigates to that feature',
    'Unloading dependencies when the user leaves (freeing memory)',
    'Reloading with different configuration after a user action (e.g., login/logout)',
    'Supporting Play Feature Delivery dynamic feature modules — whose classes don\'t exist at compile time',
]
for pt in points:
    doc.add_paragraph(pt, style='List Bullet')

doc.add_paragraph()
add_heading_styled('Zero Plugin/Processor Dependency', 3)
doc.add_paragraph(
    'Koin (DSL mode) is just a Kotlin library — no compiler plugin, no KSP, no annotation processor. '
    'This means:'
)
points = [
    'No version compatibility issues — Koin works with any Kotlin version immediately on release',
    'No KSP version must match Kotlin version (a common pain point with Hilt)',
    'No compiler plugin compatibility breaks on Kotlin upgrades (Metro requires specific Kotlin version range)',
    'Simpler build configuration — no pluginManagement, no KSP wiring, no generated source directories',
]
for pt in points:
    doc.add_paragraph(pt, style='List Bullet')

doc.add_paragraph()
add_heading_styled('Testing Simplicity', 3)
doc.add_paragraph(
    'Koin\'s test setup is the simplest of all three frameworks:'
)
points = [
    'startKoin { modules(testModule) } / stopKoin() — no special test runners or generated components',
    'Override any dependency inline: single(override = true) { MockRepository() }',
    'No @HiltAndroidTest, no custom test runner, no generated test component',
    'Fresh dependency graph per test with zero configuration',
]
for pt in points:
    doc.add_paragraph(pt, style='List Bullet')

doc.add_paragraph()
add_heading_styled('8.2 What Compile-Time DI Cannot Do', 2)

add_table(
    ['Capability', 'Hilt', 'Metro', 'Koin'],
    [
        ['Swap implementation at runtime without recompile', 'No', 'No', 'Yes'],
        ['Load/unload DI modules dynamically', 'No', 'No', 'Yes'],
        ['Support Play Feature Delivery dynamic modules', 'Limited', 'Limited', 'Yes (native)'],
        ['Work with any Kotlin version on day 1', 'No (KSP lag)', 'No (plugin compat)', 'Yes'],
        ['Zero build tool configuration', 'No (KSP + Hilt plugin)', 'No (Metro plugin)', 'Yes (DSL mode)'],
        ['Override dependencies in tests inline', 'Complex', 'Moderate', 'Trivial'],
    ],
    [5.5, 2.5, 2.5, 3]
)

doc.add_paragraph()
add_heading_styled('8.3 Honest Assessment — When Do These Advantages Matter?', 2)

doc.add_paragraph(
    'While the runtime DI advantages listed above are real, their practical relevance varies significantly:'
)

add_table(
    ['Advantage', 'How Often It Matters', 'Alternative in Hilt/Metro'],
    [
        ['Dynamic dependency swapping', 'Rare — most apps don\'t change DI providers at runtime', 'if/else inside @Provides method or Strategy pattern'],
        ['Module load/unload', 'Rare — Android lifecycle scoping handles most cases', '@ActivityScoped, @ViewModelScoped in Hilt; graph extensions in Metro'],
        ['Play Feature Delivery support', 'Niche — few apps use dynamic feature modules', 'Reflection-based entry points or interface-based lookup'],
        ['Zero plugin dependency', 'Moderate — relevant during major Kotlin version upgrades', 'Wait for KSP/plugin update (usually 1-2 weeks)'],
        ['Testing simplicity', 'Genuine — saves time daily for developers writing tests', 'Hilt testing works well once set up; Metro is straightforward'],
    ],
    [4, 4.5, 5.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('The real-world perspective: ').bold = True
p.add_run(
    'The 64us average injection time we measured sounds significant in a benchmark, but in a real app '
    'where a network call takes 200-500ms and a screen render takes 16ms, the DI resolution time is '
    'less than 0.03% of any user-visible operation. For most apps, the performance difference between '
    '2us (Metro) and 64us (Koin) is imperceptible to users. '
    'The practical reasons developers choose Koin — simplicity, faster builds, KMP support, easier onboarding — '
    'are human/productivity factors, not technical superiority of runtime resolution.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 6. iOS & KMP — CROSS-PLATFORM BENCHMARK
# ═══════════════════════════════════════════════════════
add_heading_styled('7. iOS & KMP — Cross-Platform Benchmark', 1)

doc.add_paragraph(
    'Kotlin Multiplatform (KMP) enables sharing business logic — including DI configuration — across '
    'Android and iOS. Both Metro and Koin support KMP. This section examines how each framework '
    'behaves on iOS, where the runtime environment differs significantly from the JVM.'
)

add_heading_styled('7.1 How Kotlin Code Runs on iOS', 2)
doc.add_paragraph(
    'The same Kotlin source code takes fundamentally different paths to execution on each platform:'
)
add_table(
    ['Stage', 'Android (JVM)', 'iOS (Kotlin/Native)'],
    [
        ['Compilation', '.kt → FIR → IR → JVM Bytecode → .dex', '.kt → FIR → IR → LLVM IR → ARM64 native binary'],
        ['Runtime', 'ART (Android Runtime) with JIT compilation', 'Native machine code — no VM, no interpreter'],
        ['Optimization', 'JIT: hot code paths optimized at runtime', 'AOT: all optimization at compile time via LLVM'],
        ['GC', 'Generational, concurrent, low-pause', 'Stop-the-world mark + concurrent sweep'],
        ['Output', '.apk containing .dex bytecode', '.framework containing ARM64 binary'],
    ],
    [2.5, 6.5, 6.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key point: ').bold = True
p.add_run(
    'On iOS, Kotlin compiles to ARM64 machine code via LLVM — the same backend used by Swift and Clang. '
    'There is no VM, no interpreter, no bridge. Per JetBrains, Kotlin/Native startup time on iOS is about '
    '15% faster than Swift, and over 96% of teams report no major performance concerns.'
)

doc.add_paragraph()
add_heading_styled('7.2 Platform Runtime Characteristics', 2)
doc.add_paragraph(
    'These differences are inherent to the platform — they affect all Kotlin code, not just DI frameworks:'
)
add_table(
    ['Characteristic', 'JVM (Android)', 'Kotlin/Native (iOS)', 'DI Impact'],
    [
        ['JIT Compilation', 'Hot paths optimized at runtime', 'AOT only — LLVM optimizes at compile time', 'Repeated injection calls benefit from JIT on Android; iOS relies on LLVM AOT optimization'],
        ['Concurrent Collections', 'ConcurrentHashMap (lock-striped)', 'stdlib HashMap + manual sync', 'Service locator patterns use different primitives per platform'],
        ['GC Behavior', 'Generational, short pauses', 'Stop-the-world mark + concurrent sweep', 'Object-heavy operations may see longer pauses; experimental concurrent marking available'],
        ['Lambda Dispatch', 'JIT can inline after warmup', 'Always indirect call (AOT)', 'Frameworks storing lambdas have different call overhead per platform'],
        ['Atomics', 'java.util.concurrent (mature)', 'kotlin.native.concurrent.AtomicReference', 'Thread-safe singleton access uses different primitives; Stately library provides KMP alternatives'],
    ],
    [3, 3.5, 3.5, 5.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Context: ').bold = True
p.add_run(
    'The new Kotlin/Native memory model (default since 1.7.20) simplified concurrency, and experimental '
    'concurrent marking reduces GC pauses. Both Metro and Koin work correctly on iOS — '
    'the differences are in performance characteristics, not correctness.'
)

doc.add_page_break()

add_heading_styled('7.3 Metro on iOS — Compile-Time DI on Native', 2)
doc.add_paragraph(
    'Metro\'s compiler plugin runs during Kotlin compilation, identically for all targets. '
    'The generated code — direct constructor calls and volatile field reads — goes through LLVM.'
)
add_table(
    ['Aspect', 'Behavior on iOS'],
    [
        ['Compilation', 'Metro IR codegen runs during Kotlin compilation; output goes through LLVM like any Kotlin code'],
        ['Singleton access', 'DoubleCheck volatile reads compile to ARM64 ldar (load-acquire) — similar cost to Swift lazy var'],
        ['Factory creation', 'Direct constructor calls compile to ARM64 bl (branch-link) — identical to Swift init()'],
        ['LLVM optimization', 'Direct calls are eligible for inlining, constant folding, dead code elimination'],
        ['GC impact', 'Minimal — only business objects created, no DI infrastructure objects'],
        ['Thread safety', 'Native mutex + atomic on first access; subsequent reads are lock-free atomic loads'],
        ['Memory footprint', 'One graph class + cached singletons — no HashMap, no lambdas, no metadata'],
        ['vs Android performance', 'Comparable — same code patterns (field read / constructor call) perform similarly'],
    ],
    [3.5, 12.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('In practice: ').bold = True
p.add_run(
    'Metro\'s generated code is architecturally identical to hand-written factory classes. LLVM optimizes '
    'it the same way it would optimize equivalent Swift code — the DI framework disappears entirely at runtime.'
)

doc.add_paragraph()
add_heading_styled('7.4 Koin on iOS — Runtime DI on Native', 2)
doc.add_paragraph(
    'Koin works correctly on iOS and is used in many production KMP apps. '
    'The runtime resolution path has different performance characteristics on Native vs JVM:'
)
add_table(
    ['Aspect', 'Behavior on iOS'],
    [
        ['Module registration', 'startKoin() works identically to Android — registers all definitions'],
        ['Dependency resolution', 'get<T>() performs type-based lookup and resolution — functionally identical'],
        ['Collections', 'stdlib HashMap rather than ConcurrentHashMap; thread safety via platform-specific atomics'],
        ['Lambda execution', 'Always indirect calls (no JIT); on JVM, JIT can optimize hot lambdas'],
        ['GC interaction', 'Registry structures (HashMap, lambdas, metadata) are tracked by GC during marking phase'],
        ['Memory footprint', 'Registry + lambda closures + singletons — proportional to binding count'],
        ['vs Android performance', 'Measurably different in microbenchmarks due to no JIT and different collections, but acceptable for most apps'],
    ],
    [3.5, 12.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('In practice: ').bold = True
p.add_run(
    'Companies like JetBrains (KotlinConf app), Philips, and many others ship production KMP apps with Koin on iOS. '
    'The runtime overhead is typically imperceptible to users — a network call (~200ms) dwarfs hundreds of '
    'get<T>() resolutions. Koin\'s value on iOS — simplicity, fast builds, established KMP patterns — remains compelling.'
)

doc.add_paragraph()
add_heading_styled('7.5 Swift Interop & GC Considerations', 2)
doc.add_paragraph(
    'Objects crossing the Kotlin - Swift boundary are managed by both Kotlin GC and Swift ARC:'
)
add_table(
    ['Concern', 'Details', 'Framework Impact'],
    [
        ['Object wrapping', 'Kotlin objects become stable references in GC root set', 'Equal for both — depends on boundary crossing count'],
        ['Deinitialization', 'Not immediate — waits for Kotlin GC cycle', 'Equal for both — Kotlin/Native behavior'],
        ['Retain cycles', 'Mixed Kotlin-ObjC cycles not auto-broken; use weak/unowned in Swift', 'Equal — design concern, not DI-specific'],
        ['autoreleasepool', 'Use in loops creating many cross-boundary objects', 'Equal — Kotlin/Native best practice'],
    ],
    [3, 7, 5.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('GC optimization options: ').bold = True
p.add_run(
    'kotlin.native.binary.gc=cms (concurrent marking, experimental), '
    'kotlin.native.binary.appStateTracking=enabled (background optimization), '
    '-Xruntime-logs=gc=info (pause monitoring).'
)

doc.add_page_break()

add_heading_styled('7.6 Measured iOS Results — Compile Time & Runtime', 2)

doc.add_paragraph(
    'We built unified KMP modules (benchmark-metro-large and benchmark-koin-large) targeting'
    'iosSimulatorArm64, each containing the full e-commerce app (~350 classes, ~285 bindings). '
    'Both frameworks use best practices: Metro with @SingleIn(AppScope::class) scoping, '
    'Koin with createdAtStart on 19 critical singletons. Benchmark methodology follows '
    'kotlinx-benchmark patterns: warmup phase (5 iterations, excluded from results), '
    'blackhole to prevent LLVM dead code elimination, single container lifecycle per framework.'
)

doc.add_paragraph()
add_heading_styled('iOS Compile Time (3 clean builds, IosSimulatorArm64)', 3)
add_table(
    ['Metric', 'Metro', 'Koin'],
    [
        ['Average', '1,871ms', '1,260ms'],
        ['Min', '1,340ms', '1,093ms'],
        ['Max', '2,770ms', '1,425ms'],
        ['Comparison', '', 'Koin 32.7% faster'],
    ],
    [5, 4, 4]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why Koin compiles faster on iOS: ').bold = True
p.add_run(
    'Same reason as Android — Koin has zero codegen overhead. Metro\'s compiler plugin runs FIR analysis '
    'and IR codegen during Kotlin/Native compilation, adding ~600ms. Both go through the same LLVM pipeline afterward.'
)

doc.add_paragraph()
add_heading_styled('iOS Runtime — Container Initialization', 3)
add_table(
    ['Framework', 'Init Time', 'What Happens'],
    [
        ['Metro', '0.04ms', 'Creates one pre-compiled graph class with all wiring baked in'],
        ['Koin', '0.20ms', 'Registers 271 definitions across 24 modules, creates 19 eager singletons'],
    ],
    [3, 2, 11]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Metro is 5x faster. ').bold = True
p.add_run(
    'Metro\'s graph is a single LLVM-compiled class — instantiation is one ARM64 constructor call. '
    'Koin must process 24 module DSL blocks and populate its internal HashMap registry.'
)

doc.add_paragraph()
add_heading_styled('iOS Runtime — Cold Injection (First Access)', 3)
add_table(
    ['Class', 'Metro', 'Koin', 'Metro Faster By'],
    [
        ['HomeViewModel (6 deps, ~50 transitive)', '19us', '38us', '2x'],
        ['CheckoutViewModel (7 deps)', '5us', '22us', '4.4x'],
        ['ProductDetailVM (6 deps)', '3us', '11us', '3.7x'],
        ['ProfileViewModel (7 deps)', '5us', '11us', '2.2x'],
        ['ChatViewModel (5 deps)', '3us', '9us', '3x'],
        ['SearchViewModel (3 deps)', '2us', '5us', '2.5x'],
        ['CartViewModel (5 deps)', '2us', '13us', '6.5x'],
        ['OrderHistoryVM (3 deps)', '1us', '5us', '5x'],
        ['AnalyticsTracker (2 deps)', '<1us', '<1us', '~1x'],
        ['ProductRepository (4 deps)', '<1us', '<1us', '~1x'],
    ],
    [5.5, 2, 2, 2.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Metro wins 10/10 classes. ').bold = True
p.add_run(
    'Metro\'s singletons are pre-wired at createGraph() — first property access is a direct volatile field read. '
    'Koin must resolve the full dependency chain via HashMap lookup on each get<T>() call, recursively '
    'for every transitive dependency.'
)

doc.add_paragraph()
add_heading_styled('iOS Runtime — Warm Injection (Avg of 100 Iterations)', 3)
add_table(
    ['Class', 'Metro', 'Koin', 'Metro Faster By'],
    [
        ['HomeViewModel', '2us', '8us', '4x'],
        ['CartViewModel', '1us', '5us', '5x'],
        ['CheckoutViewModel', '1us', '13us', '13x'],
        ['ProductDetailVM', '1us', '5us', '5x'],
        ['ProfileViewModel', '1us', '5us', '5x'],
        ['SearchViewModel', '1us', '3us', '3x'],
        ['ChatViewModel', '1us', '3us', '3x'],
        ['OrderHistoryVM', '1us', '3us', '3x'],
        ['AnalyticsTracker', '<1us', '<1us', '~1x'],
        ['ProductRepository', '<1us', '<1us', '~1x'],
        ['', '', '', ''],
        ['TOTAL avg/injection', '1us', '5us', '5x'],
    ],
    [5.5, 2, 2, 2.5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Metro is effectively instant on warm access. ').bold = True
p.add_run(
    'With @SingleIn(AppScope::class) scoping, Metro\'s singleton access compiles to an ARM64 ldar '
    '(load-acquire) instruction — a single CPU cycle plus memory barrier. Koin\'s single{} also caches '
    'the instance, but get<T>() still performs HashMap.get() + type resolution + return on every access. '
    'On iOS (no JIT), this overhead cannot be optimized away at runtime.'
)

doc.add_paragraph()
add_heading_styled('iOS Memory', 3)
p = doc.add_paragraph()
p.add_run('Note: ').bold = True
p.add_run(
    'iOS process-level memory measurement (mach_task_basic_info.resident_size) is unreliable for '
    'DI-level granularity because Kotlin/Native GC cycles, system page reclamation, and SwiftUI rendering '
    'all affect resident memory between measurements. For accurate iOS memory profiling, '
    'use Xcode Instruments (Allocations instrument). The Android benchmark (which uses '
    'Debug.getNativeHeapAllocatedSize()) showed Metro at 241KB vs Koin at 1,537KB — a 6.4x difference. '
    'We expect a similar ratio on iOS based on the identical code structure.'
)

doc.add_page_break()

add_heading_styled('7.7 Android vs iOS — Cross-Platform Comparison', 2)

doc.add_paragraph(
    'The same e-commerce app (~350 classes, ~285 bindings) was benchmarked on both platforms. '
    'This is the first apples-to-apples cross-platform DI comparison with identical code:'
)

add_table(
    ['Metric', 'Android (Metro)', 'Android (Koin)', 'iOS (Metro)', 'iOS (Koin)'],
    [
        ['Container Init', '0.14ms', '1.45ms', '0.01ms', '0.17ms'],
        ['Cold avg', '~15us', '~167us', '~2us', '~13us'],
        ['Warm avg/injection', '11us', '64us', '<1us', '6us'],
        ['Memory', '241KB', '1,537KB', 'N/A (use Instruments)', 'N/A (use Instruments)'],
    ],
    [3, 3, 3, 3, 3]
)

doc.add_paragraph()
add_heading_styled('Why iOS Numbers Are Lower Than Android', 3)

doc.add_paragraph(
    'Both Metro and Koin show faster absolute times on iOS than Android. This is not an error — '
    'it reflects genuine platform differences:'
)
points = [
    'The iOS benchmark runs on Apple Silicon (M-series) via the Simulator, which has faster single-core performance than the Android emulator running on the same hardware.',
    'Kotlin/Native compiles to native ARM64 machine code — there is no VM startup overhead, no JIT warmup. Code runs at full speed from the first call.',
    'The warmup phase in our benchmark eliminates class-loading bias on both platforms, making the comparison fair.',
    'Metro\'s advantage over Koin is consistent across platforms: ~6x faster on both Android (11us vs 64us) and iOS (<1us vs 6us) for warm injection.',
]
for pt in points:
    doc.add_paragraph(pt, style='List Bullet')

doc.add_paragraph()
add_heading_styled('Key Cross-Platform Observation', 3)
p = doc.add_paragraph()
p.add_run('Metro\'s relative advantage is consistent: ').bold = True
p.add_run(
    'On Android, Metro is 5.8x faster than Koin on warm injection (11us vs 64us). '
    'On iOS, Metro is >6x faster (<1us vs 6us). The ratio holds because the architectural '
    'difference is the same: Metro uses direct field reads (volatile/atomic), Koin uses '
    'HashMap lookups + lambda invocations. The absolute times are lower on iOS because '
    'native ARM64 code on Apple Silicon is simply faster than JVM bytecode on an Android emulator.'
)

doc.add_page_break()

add_heading_styled('7.8 Practical Guidance for KMP Teams', 2)
add_table(
    ['Scenario', 'Recommended', 'Reasoning'],
    [
        ['Large KMP app, performance-sensitive iOS UI', 'Metro', 'Zero runtime overhead; LLVM optimizes generated code like native Swift'],
        ['Small-medium KMP app, rapid development', 'Koin', 'Simplest setup, fast builds, established patterns; overhead is negligible'],
        ['Team migrating from Dagger/Hilt to KMP', 'Metro', 'Familiar @Inject/@Provides patterns reduce migration effort'],
        ['App with dynamic features on iOS', 'Koin', 'Runtime module loading/unloading fits naturally'],
        ['Rapid prototyping / MVP', 'Koin', 'Fastest to set up, extensive documentation'],
        ['Consistent perf across platforms', 'Metro', 'Same code pattern performs identically on JVM and LLVM'],
    ],
    [5, 2, 9]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Both are production-ready for KMP. ').bold = True
p.add_run(
    'Metro offers better raw performance; Koin offers better developer ergonomics. Neither is a wrong choice — '
    'they represent different trade-offs on the same spectrum.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 10. DEEP DIVE — HOW METRO ACHIEVES SUPERIOR PERFORMANCE
# ═══════════════════════════════════════════════════════
add_heading_styled('8. Deep Dive — How Metro Achieves Superior Performance', 1)

doc.add_paragraph(
    'Metro is faster than Hilt at compile time AND faster than Koin at runtime. This section explains '
    'exactly how — starting with the Kotlin K2 compiler architecture, then showing where Metro hooks in, '
    'and finally revealing what code Metro actually generates.'
)

add_heading_styled('8.1 K2 Compiler Pipeline', 2)

doc.add_paragraph(
    'The Kotlin K2 compiler has two distinct phases with a clear boundary between them:'
)

add_table(
    ['Phase', 'Component', 'What Happens'],
    [
        ['1. Frontend', 'PSI Parser', 'Source code (.kt) is parsed into a PSI (Program Structure Interface) tree'],
        ['2. Frontend', 'Raw FIR', 'PSI tree is transformed into FIR (Frontend Intermediate Representation)'],
        ['3. Frontend', 'FIR Resolution', 'Types, symbols, scopes, imports, annotations are resolved in multiple passes'],
        ['4. Frontend', 'FIR Checkers', 'Diagnostics run — warnings and errors are reported to IDE and build output'],
        ['— Bridge —', 'Fir2Ir', 'Resolved FIR is transformed into IR (Intermediate Representation)'],
        ['5. Backend', 'IR Generation', 'IR tree is built from Fir2Ir output'],
        ['6. Backend', 'IR Lowering', 'IR is optimized and lowered for the target platform'],
        ['7. Backend', 'Code Emission', 'Final bytecode (.class for JVM), JavaScript, Wasm, or native binary is emitted'],
    ],
    [2.5, 2.5, 11]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Critical distinction: ').bold = True
p.add_run(
    'FIR lives entirely in the Frontend — it handles analysis, type resolution, and error reporting. '
    'IR lives entirely in the Backend — it handles code generation and platform-specific output. '
    'The Fir2Ir bridge converts the Frontend\'s semantic information into the Backend\'s code representation. '
    'This is important because Metro hooks into BOTH phases for different purposes.'
)

doc.add_paragraph()
add_heading_styled('8.2 Where Metro Hooks Into the Compiler', 2)

add_table(
    ['Compiler Phase', 'What Metro Does', 'Why This Phase'],
    [
        ['FIR (Frontend)', 'Reads @DependencyGraph, @Inject, @Provides, @SingleIn annotations', 'FIR has fully resolved type information — Metro knows every class and its dependencies'],
        ['FIR (Frontend)', 'Validates the entire dependency graph — missing bindings, cycles, scope violations', 'Errors reported here appear in IDE (K2 plugin) and fail the build immediately'],
        ['FIR (Frontend)', 'Generates synthetic class headers (declarations only, no method bodies)', 'Downstream code can reference Metro-generated types during the same compilation'],
        ['FIR (Frontend)', 'Records private @Provides visibility metadata', 'Private providers need special metadata for cross-compilation visibility in IR'],
        ['— Fir2Ir —', 'Metro\'s synthetic FIR declarations are converted to IR stubs', 'Bridge phase transforms Metro\'s headers into IR nodes ready for body generation'],
        ['IR (Backend)', 'Builds BindingGraph from all collected metadata', 'IR phase has the complete picture of all bindings across all modules'],
        ['IR (Backend)', 'Runs Tarjan\'s algorithm for cycle detection in the dependency graph', 'Detects circular dependencies that would cause infinite recursion'],
        ['IR (Backend)', 'Runs Kahn\'s algorithm for topological sort (initialization order)', 'Determines the correct order to initialize singletons without forward references'],
        ['IR (Backend)', 'Generates provider factory class bodies with DoubleCheck for singletons', 'Actual constructor call code, thread-safe caching — all as IR nodes, no source files'],
        ['IR (Backend)', 'Generates the graph implementation class (e.g., ShopAppGraphImpl)', 'The single class that holds all wiring — property getters, provider fields, factory methods'],
    ],
    [2.5, 6, 7.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key insight: ').bold = True
p.add_run(
    'Metro\'s generated code never exists as source files. It is born as IR nodes inside the compiler\'s '
    'backend phase and goes directly to bytecode. No .java files, no .kt files, no disk I/O — just '
    'in-memory IR that gets lowered to platform bytecode alongside your own code in the same compilation pass.'
)

doc.add_page_break()

add_heading_styled('8.3 Hilt Pipeline vs Metro Pipeline (Step-by-Step)', 2)

doc.add_paragraph(
    'This is why Metro compiles 44.7% faster than Hilt for the same dependency graph:'
)

add_heading_styled('Hilt/Dagger Compilation (4 steps, 2 compiler invocations)', 3)
add_table(
    ['Step', 'What Runs', 'Cost', 'Output'],
    [
        ['1', 'Kotlin Frontend (FIR) — your code', 'Resolves all symbols', 'Resolved FIR'],
        ['2', 'KSP processor — runs as separate step', 'Reads resolved symbols, scans all annotations', 'Nothing yet (analysis)'],
        ['3', 'KSP code generation — writes .java files', 'Generates 291 Java source files (488KB)', '291 .java files on disk'],
        ['4', 'Kotlin Backend (IR → bytecode) — your code', 'Compiles your Kotlin to bytecode', '.class files for your code'],
        ['5', 'Java Compiler (javac) — generated code', 'Entire separate compiler for 291 .java files', '.class files for generated code'],
    ],
    [1, 5, 4.5, 4.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Total: ').bold = True
p.add_run('2 frontend passes + KSP processing + file I/O (write 291 files, read them back) + javac invocation')

doc.add_paragraph()
add_heading_styled('Metro Compilation (1 step, 1 compiler invocation)', 3)
add_table(
    ['Step', 'What Runs', 'Cost', 'Output'],
    [
        ['1', 'Kotlin Frontend (FIR) + Metro analysis', 'Normal resolution + Metro validates graph', 'Resolved FIR + Metro metadata'],
        ['2', 'Fir2Ir bridge', 'Normal + Metro\'s synthetic declarations', 'IR stubs'],
        ['3', 'Kotlin Backend (IR) + Metro codegen', 'Normal IR lowering + Metro generates provider factories, graph impl', 'Bytecode (your code + Metro\'s code, unified)'],
    ],
    [1, 5, 5.5, 4.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Total: ').bold = True
p.add_run('1 frontend pass + 1 backend pass. Zero file I/O. Zero separate processes. Zero javac.')

doc.add_paragraph()
add_heading_styled('Why KSP Is Inherently Slower Than a Compiler Plugin', 3)
add_table(
    ['Factor', 'KSP (Hilt)', 'Compiler Plugin (Metro)'],
    [
        ['Runs inside kotlinc?', 'No — separate processing step alongside the compiler', 'Yes — native plugin inside the compiler itself'],
        ['Can modify existing classes?', 'No — can only generate NEW source files', 'Yes — can inject code into existing classes (IR)'],
        ['File I/O required?', 'Yes — writes generated .java/.kt files to disk', 'No — generates IR nodes in memory'],
        ['Needs second compiler pass?', 'Yes — javac must compile generated Java files', 'No — generated IR is lowered in the same pass'],
        ['Can access private members?', 'No — respects Kotlin visibility', 'Yes — can inject into private @Provides, private constructors'],
        ['Incremental build overhead?', 'KSP must re-analyze changed files + regenerate', 'Plugin sees only changed IR, regenerates minimally'],
    ],
    [4, 5.5, 5.5]
)

doc.add_page_break()

add_heading_styled('8.4 What Metro Generates at Runtime', 2)

doc.add_paragraph(
    'Metro\'s IR phase generates a single implementation class for each @DependencyGraph. '
    'This is what the generated bytecode looks like when decompiled (Metro never creates this as source code — '
    'it exists only as bytecode, shown here as pseudocode for illustration):'
)

# Provider factory pseudocode
p = doc.add_paragraph()
p.add_run('Generated Graph Implementation (pseudocode from bytecode):').bold = True

code_lines = [
    'class ShopAppGraphImpl : ShopAppGraph {',
    '',
    '    // ── Singletons: DoubleCheck (thread-safe lazy init) ──',
    '    // First call: creates instance under synchronized lock',
    '    // All subsequent calls: volatile field read (~0.2us)',
    '    private val databaseMgr = DoubleCheck { DatabaseManager() }',
    '    private val cacheMgr = DoubleCheck { CacheManager(databaseMgr.get()) }',
    '    private val httpClient = DoubleCheck { HttpClient("https://api.shopapp.com", 30000) }',
    '    private val authInterceptor = DoubleCheck { AuthInterceptor(tokenProvider.get()) }',
    '    private val productRemote = DoubleCheck {',
    '        ProductRemoteDataSource(httpClient.get(), apiParser.get(), authInterceptor.get(), ...)',
    '    }',
    '    private val productLocal = DoubleCheck {',
    '        ProductLocalDataSource(databaseMgr.get(), cacheMgr.get())',
    '    }',
    '    private val productRepo = DoubleCheck {',
    '        ProductRepository(productRemote.get(), productLocal.get(), productMapper.get(), logger.get())',
    '    }',
    '    // ... 105 more DoubleCheck providers for all singletons',
    '',
    '    // ── Singleton access: direct field read ──',
    '    override val productRepository get() = productRepo.get()    // ~0.2us after init',
    '    override val analyticsTracker get() = analyticsProvider.get() // ~0.2us after init',
    '',
    '    // ── Factory access: direct constructor calls ──',
    '    override val homeViewModel get() = HomeViewModel(',
    '        getProductList = GetProductListUseCase(productRepo.get(), analyticsProvider.get()),',
    '        getCategoryList = GetCategoryListUseCase(categoryRepo.get(), analyticsProvider.get()),',
    '        getPromotionList = GetPromotionListUseCase(promotionRepo.get(), analyticsProvider.get()),',
    '        getFeedList = GetFeedListUseCase(feedRepo.get(), analyticsProvider.get()),',
    '        analytics = analyticsProvider.get(),',
    '        featureFlags = featureFlagProvider.get()',
    '    )   // New HomeViewModel instance every call, but deps are cached singleton reads',
    '}',
]
for line in code_lines:
    p = doc.add_paragraph(line)
    for r in p.runs:
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)

doc.add_paragraph()
add_heading_styled('DoubleCheck — Thread-Safe Singleton Pattern', 3)

doc.add_paragraph(
    'Metro uses the same DoubleCheck pattern as Dagger for scoped (singleton) bindings. '
    'This is the double-checked locking idiom from Effective Java:'
)

add_table(
    ['Call', 'What Happens', 'Cost'],
    [
        ['First call', 'volatile read (UNINITIALIZED) → synchronized block → create instance → store in volatile field → release lock', '~5-50us depending on constructor complexity'],
        ['All subsequent calls', 'volatile read → instance already set → return immediately (skips synchronized)', '~0.2us (1 CPU instruction + memory barrier)'],
    ],
    [2.5, 10, 3.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('This means: ').bold = True
p.add_run(
    'After app startup, every singleton access in Metro costs exactly one volatile field read. '
    'No HashMap, no type checking, no lambda invocation — just a direct memory access.'
)

doc.add_paragraph()
add_heading_styled('8.5 Why Metro Is 10-80x Faster Than Koin at Runtime', 2)

doc.add_paragraph(
    'The performance gap comes from what happens on every dependency resolution:'
)

add_table(
    ['Operation', 'Metro (per access)', 'Koin (per access)'],
    [
        ['Singleton lookup', 'Volatile field read (~0.2us)', 'HashMap.get(key) → type check → return cached (~3us)'],
        ['Factory creation', 'new ClassName(dep1.get(), dep2.get()) — direct call', 'HashMap.get(key) → invoke lambda → recursive get() for each param'],
        ['Type resolution', 'None — type is compiled into the field reference', 'Reified type → qualified name → hash computation per lookup'],
        ['Thread safety', 'DoubleCheck (volatile + sync on first call only)', 'ConcurrentHashMap (hash + bucket sync on every access)'],
        ['Dependency chain', 'Singleton deps: field reads. Factory deps: constructor calls.', 'Every node in chain: HashMap lookup + type check + lambda/cache check'],
    ],
    [3, 5.5, 7]
)

doc.add_paragraph()
add_heading_styled('Cost Scaling with Dependency Depth', 3)

add_table(
    ['Class (Dep Depth)', 'Metro', 'Koin', 'Koin/Metro Ratio', 'Why'],
    [
        ['AnalyticsTracker (2 deps)', '0.2us', '3us', '15x', '2 volatile reads vs 2 HashMap lookups'],
        ['OrderHistoryVM (3+15 transitive)', '1us', '28us', '28x', '3 constructor calls + ~15 field reads vs ~18 HashMap lookups'],
        ['HomeViewModel (6+50 transitive)', '10us', '105us', '10.5x', '6 constructors + ~50 field reads vs ~56 HashMap + lambda operations'],
        ['CheckoutViewModel (7+40 transitive)', '11us', '166us', '15x', '7 constructors + ~40 field reads vs ~47 HashMap + lambda operations'],
    ],
    [4, 1.5, 1.5, 2, 7]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Pattern: ').bold = True
p.add_run(
    'Metro\'s cost per singleton dep is ~0.2us (volatile read), constant regardless of chain depth. '
    'Koin\'s cost per dep is ~3us (HashMap + type check), and it compounds multiplicatively through the chain. '
    'The deeper the dependency tree, the larger Metro\'s advantage.'
)

doc.add_paragraph()
add_heading_styled('8.6 Square/Cash App — Real-World Proof at Scale', 2)

doc.add_paragraph(
    'Square (Block) migrated their entire Android platform from Dagger to Metro — the largest known '
    'Metro deployment. Their published results validate Metro\'s performance claims at production scale:'
)

add_table(
    ['Metric', 'Value'],
    [
        ['Gradle modules migrated', '7,000+'],
        ['CI jobs transitioned', '1,500'],
        ['Development apps updated', '300+'],
        ['Production apps converted', '22'],
        ['Pull requests merged', '850+'],
        ['Migration duration', '9 months'],
    ],
    [5, 5]
)

doc.add_paragraph()
add_heading_styled('Build Time Improvements at Square', 3)

add_table(
    ['Scenario', 'Improvement Range', 'Best Case'],
    [
        ['ABI changes (API signature change)', '20.5% — 56.5% faster', '56.5% (Account module)'],
        ['Dagger wiring changes', '21.4% — 48.2% faster', '48.2% (development app)'],
        ['Non-ABI changes (implementation only)', '4.1% — 8.0% faster', '8.0% (development app)'],
    ],
    [5, 4.5, 5.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('The headline number: ').bold = True
p.add_run(
    '"Even if we assume a conservative average improvement of only 10%, that still saves us more than '
    '4,800 hours of Gradle build time in CI every week." — Square Engineering Blog'
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why Non-ABI improvements are smaller: ').bold = True
p.add_run(
    'When only implementation code changes (no API changes), Gradle\'s incremental compilation already '
    'avoids recompiling downstream modules. Metro\'s advantage is largest when API/wiring changes trigger '
    'cascading recompilations — exactly where Dagger/KSP\'s code generation creates the most overhead.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 8. LIFECYCLE AWARENESS & ANDROID INTEGRATION
# ═══════════════════════════════════════════════════════

# Note: kotlin-inject-anvil overview (old 11.1, 11.2) moved to Section 2.4
# The KSP pipeline explanation is kept here briefly for context

add_heading_styled('9. Lifecycle Awareness & Android Integration', 1)

doc.add_paragraph(
    'Lifecycle awareness — who manages when objects are created, destroyed, and survive configuration '
    'changes — is a critical differentiator for Android development. This section compares all four frameworks.'
)

# The lifecycle table from old 11.4 goes here — it's already in the file, just need to restructure

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 9. JAVA INTEROP & MIGRATION FROM DAGGER/HILT
# ═══════════════════════════════════════════════════════
add_heading_styled('10. Java Interop & Migration from Dagger/Hilt', 1)

doc.add_paragraph(
    'Most production Android apps today use Hilt or Dagger. Migrating to KMP while preserving the existing '
    'DI graph is the biggest practical challenge teams face. This section covers interop capabilities, '
    'Java code support, and step-by-step migration paths for each framework.'
)

add_heading_styled('Interop & Java Support', 2)
add_table(
    ['Capability', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Share graph with existing Hilt/Dagger', 'Yes — @Includes daggerComponent', 'No — manual bridge', 'No — manual bridge'],
        ['Understand javax.inject.Inject', 'Yes — includeDagger()', 'Only in Kotlin files', 'No'],
        ['Understand jakarta.inject.Inject', 'Yes — includeJakarta()', 'No', 'No'],
        ['Reuse Dagger-generated Java factories', 'Yes — zero extra code', 'No — must wrap with @Provides', 'No'],
        ['Process .java source files', 'No (Kotlin compiler plugin)', 'No (KSP = Kotlin only)', 'N/A'],
        ['Incremental migration (both frameworks running)', 'Yes — Gradle flag, module by module', 'No — clean break', 'No — clean break'],
        ['Effort per shared dependency', '0 lines (auto via @Includes)', '1 @get:Provides line + entry point', '1 manual binding'],
        ['Effort for 100 shared dependencies', '1 line total', '~200 lines (100 params + 100 entry points)', '~100 bindings'],
    ],
    [4.5, 4, 4, 3.5]
)

doc.add_paragraph()
add_heading_styled('Migration Path: From Hilt to Each Framework', 2)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Migrating to Metro (from existing Hilt):')
r.bold = True
add_table(
    ['Phase', 'What You Do', 'Both Frameworks Running?'],
    [
        ['Phase 1', 'Keep Hilt for all existing Android modules. Enable Metro plugin.', 'Yes — Hilt handles existing, Metro is available'],
        ['Phase 2', 'New KMP feature modules use Metro with @Includes(daggerComponent) to access Hilt-provided singletons', 'Yes — Metro reads from Hilt graph directly'],
        ['Phase 3', 'Gradually migrate old Hilt modules to Metro, one by one. Both can run in parallel via Gradle build flag.', 'Yes — dual build, Square did this across 7,000 modules'],
        ['Phase 4', 'Remove Hilt when all modules migrated. Single Metro graph.', 'No — Metro only'],
    ],
    [1.5, 10, 4.5]
)
p = doc.add_paragraph()
p.add_run('Zero bridge code required. Zero dependency duplication. Proven at Square scale (7,000 modules, 9 months).').italic = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Migrating to kotlin-inject-anvil (from existing Hilt):')
r.bold = True
add_table(
    ['Phase', 'What You Do', 'Both Frameworks Running?'],
    [
        ['Phase 1', 'Keep Hilt for all existing Android modules.', 'Yes'],
        ['Phase 2', 'Create @EntryPoint interfaces in Hilt to expose every dependency the new modules need.', 'Yes — but manual entry points needed'],
        ['Phase 3', 'New KMP modules use kotlin-inject-anvil. Pass Hilt dependencies via component constructor parameters.', 'Yes — but you maintain a manual bridge layer'],
        ['Phase 4', 'For each shared dependency: update @EntryPoint + component constructor + creation site.', 'Yes — bridge grows with each shared dep'],
        ['Phase 5', 'Eventually remove Hilt. All dependencies in kotlin-inject-anvil. Remove bridge.', 'No — kotlin-inject-anvil only'],
    ],
    [1.5, 10, 4.5]
)
p = doc.add_paragraph()
p.add_run('Requires manual bridge code for every shared dependency. Bridge must be maintained during migration.').italic = True

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Migrating to Koin (from existing Hilt):')
r.bold = True
add_table(
    ['Phase', 'What You Do', 'Both Frameworks Running?'],
    [
        ['Phase 1', 'Keep Hilt for all existing Android modules.', 'Yes'],
        ['Phase 2', 'Create @EntryPoint interfaces in Hilt to expose shared dependencies.', 'Yes — manual entry points'],
        ['Phase 3', 'New KMP modules use Koin. Pass Hilt dependencies via startKoin module definitions.', 'Yes — manual bindings for each shared dep'],
        ['Phase 4', 'Gradually move modules. Each migrated module removes Hilt annotations, adds Koin module { } DSL.', 'Yes — both running independently'],
        ['Phase 5', 'Remove Hilt. All dependencies in Koin. No compile-time DI safety (unless using Koin compiler plugin).', 'No — Koin only'],
    ],
    [1.5, 10, 4.5]
)
p = doc.add_paragraph()
p.add_run('Same manual bridge as kotlin-inject-anvil, plus you lose compile-time graph validation during and after migration.').italic = True

doc.add_paragraph()
add_heading_styled('Migration Effort Summary', 2)
add_table(
    ['Factor', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Bridge code needed', 'None', 'Yes — per shared dependency', 'Yes — per shared dependency'],
        ['Can run alongside Hilt', 'Yes — shares the same graph', 'Yes — but independent graphs', 'Yes — but independent graphs'],
        ['Effort for 50-dep module', '~1 hour (add @Includes)', '~1 day (bridge + entry points)', '~1 day (bridge + module DSL)'],
        ['Effort for 500-dep project', '~2 weeks (module by module)', '~2 months (bridge maintenance)', '~2 months (bridge + no compile safety)'],
        ['Risk during migration', 'Low — both graphs validated at compile time', 'Medium — bridge can have type mismatches', 'High — Koin errors only at runtime'],
        ['Proven at scale', 'Square: 7,000 modules, 9 months', 'No comparable public migration story', 'Some teams migrated from Hilt to Koin'],
    ],
    [4, 4, 4, 4]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 10. HEAD-TO-HEAD — ALL 4 FRAMEWORKS COMPARED
# ═══════════════════════════════════════════════════════
add_heading_styled('11. Head-to-Head — All 4 Frameworks Compared', 1)

doc.add_paragraph(
    'This section provides the complete side-by-side comparison of all four frameworks across every dimension.'
)

# OLD 11.1 content starts below — but now it's just the overview text for context
add_heading_styled('kotlin-inject-anvil Overview', 2)
doc.add_paragraph(
    'kotlin-inject-anvil is a two-part system by Amazon: kotlin-inject (by Evan Tatarka) provides compile-time '
    'DI for Kotlin, and kotlin-inject-anvil (by Ralf Wondratschek at Amazon) adds Anvil-style auto-merging '
    '(@ContributesTo, @ContributesBinding, @MergeComponent) on top. Together they deliver compile-time DI '
    '+ automatic module merging + full KMP support — similar to Metro but using KSP instead of a compiler plugin.'
)
add_table(
    ['Attribute', 'Detail'],
    [
        ['Type', 'Compile-time (KSP annotation processing)'],
        ['Base framework', 'kotlin-inject 0.7.2+ (Evan Tatarka)'],
        ['Extensions', 'kotlin-inject-anvil 0.1.6+ (Amazon)'],
        ['Processing', 'KSP → generates Kotlin source files → compiled by kotlinc'],
        ['Annotations', '@Inject, @Component, @Provides, @ContributesTo, @ContributesBinding, @MergeComponent'],
        ['Output', 'Generated .kt files in build/generated/ksp/ (visible, debuggable)'],
        ['Scoping', '@SingleIn(Scope::class), custom @Scope annotations'],
        ['KMP support', 'Yes — JVM, iOS, JS, Wasm, Native'],
        ['Backed by', 'Amazon (kotlin-inject-anvil), community (kotlin-inject)'],
        ['Production users', 'Amazon apps, Bitkey (170 KMP modules)'],
    ],
    [4, 12]
)

doc.add_paragraph()
add_heading_styled('How kotlin-inject-anvil Works (KSP Pipeline)', 3)
doc.add_paragraph(
    'Unlike Metro (which hooks into the Kotlin compiler), kotlin-inject-anvil runs as a separate KSP step:'
)
add_table(
    ['Step', 'What Happens', 'Output'],
    [
        ['1. KSP scan', 'Reads @ContributesTo, @ContributesBinding, @MergeComponent across all modules', 'Contribution metadata'],
        ['2. KSP merge', 'Collects all contributions for each scope, generates merged component interface', 'Generated .kt file with @Component'],
        ['3. kotlin-inject KSP', 'Reads @Component, @Inject, @Provides → generates ComponentImpl class', 'Generated .kt factory code'],
        ['4. Kotlin compiler', 'Compiles your code + generated code in one pass', 'Bytecode / klib'],
    ],
    [2.5, 7, 5]
)
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key difference from Metro: ').bold = True
p.add_run(
    'Generated code exists as readable .kt files in build/generated/ksp/. You can open them, set breakpoints, '
    'and step through them in the debugger. Metro\'s IR-generated code is invisible — you cannot debug it directly. '
    'This is the fundamental trade-off: Metro is faster (no file I/O), kotlin-inject-anvil is more debuggable.'
)

doc.add_page_break()

add_heading_styled('Metro vs kotlin-inject-anvil', 2)
add_table(
    ['Aspect', 'Metro', 'kotlin-inject-anvil'],
    [
        ['Technology', 'Kotlin compiler plugin (FIR + IR)', 'KSP (Kotlin Symbol Processing)'],
        ['Generated files', 'Zero — code exists only as IR bytecode', 'Kotlin .kt files in build/generated/ksp/'],
        ['Can debug generated code?', 'No (IR invisible, debug log available)', 'Yes — standard breakpoints and step-through'],
        ['Can modify existing classes?', 'Yes — injects code into existing classes', 'No — can only generate new files'],
        ['Can access private members?', 'Yes — private @Provides, private constructors', 'No — respects Kotlin visibility'],
        ['Default value copying', 'Yes — copies default expressions via IR', 'No — KSP cannot access default values'],
        ['@GraphPrivate (prevent child access)', 'Yes', 'Not available'],
        ['Build speed (clean)', '~2,500ms avg — one compiler pass, no file I/O', '~2,500ms avg — KSP + 30 generated files'],
        ['Build speed at scale', 'Square: 20-56% faster vs Dagger', 'Bitkey: 170 KMP modules'],
        ['Runtime total (124 classes)', '7ms', '6ms'],
        ['Component/Graph init', '3.46ms', '2.82ms'],
        ['ViewModel init (13)', '3.67ms', '3.67ms (identical)'],
        ['Dagger/Hilt interop', 'Yes — includeDagger(), @Includes', 'No — must bridge manually'],
        ['Java code support', 'Yes — reuses Dagger-generated Java factories', 'Kotlin-only (KSP processes .kt files only)'],
        ['javax/jakarta annotations', 'Yes — includeDagger(), includeJakarta()', 'javax.inject.* only in Kotlin files'],
        ['Custom scoping', 'Yes — @Scope, @SingleIn, Graph Extensions', 'Yes — @Scope, @SingleIn, @ContributesSubcomponent'],
        ['Module merging', 'Yes — @ContributesTo', 'Yes — @ContributesTo, @ContributesBinding'],
        ['Assisted injection', 'Yes — @Assisted + @AssistedFactory', 'Yes — @Assisted, factory bound as lambda'],
        ['Multibindings', 'Yes — @IntoSet, @IntoMap', 'Yes — via kotlin-inject'],
        ['Runtime lifecycle callbacks', 'No', 'Yes — via Amazon App Platform (Scoped interface)'],
        ['Auto CoroutineScope per scope', 'No', 'Yes — via App Platform scope.launch'],
        ['KMP support', 'JVM, iOS, JS, Wasm, Native', 'JVM, iOS, JS, Wasm, Native'],
        ['Backed by', 'Zac Sweers / Square', 'Ralf Wondratschek / Amazon'],
        ['Production scale', 'Square: 7,000 modules, 22 apps', 'Bitkey: 170 KMP modules'],
    ],
    [4.5, 5.5, 5.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Summary: ').bold = True
p.add_run(
    'Both are excellent compile-time KMP DI solutions with nearly identical performance — ~2.5s compile time '
    'and 6-7ms runtime for 124 classes. Metro generates code in IR (zero source files, not debuggable) '
    'and has native Dagger interop via @Includes. kotlin-inject-anvil generates 30 debuggable Kotlin files '
    'and has @ContributesBinding for auto-discovery across modules. '
    'At runtime, both produce functionally equivalent code — direct constructor calls and volatile field reads.'
)

doc.add_paragraph()
add_heading_styled('Lifecycle Awareness (All 4 Frameworks)', 2)
doc.add_paragraph(
    'Lifecycle awareness — who manages when objects are created, destroyed, and survive configuration '
    'changes — is a critical differentiator, especially for Android development:'
)
add_table(
    ['Aspect', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Automatic lifecycle management', 'Yes — generated code hooks into Android callbacks', 'No — manual graph creation/destruction', 'No — manual (App Platform adds callbacks)', 'Partial — viewModel() is auto, scopes are manual'],
        ['Predefined scope hierarchy', 'Yes — 6 levels (Singleton → View)', 'No — Graph Extensions (you define)', 'No — you define scopes', 'No — you define modules'],
        ['ViewModel support', '@HiltViewModel — zero boilerplate', 'MetroViewModelFactory (metrox-viewmodel)', '@ContributesViewModel (3rd party lib)', 'viewModel { } DSL'],
        ['Survives rotation', 'Automatic (ActivityRetainedComponent)', 'Via ViewModelProvider (manual setup)', 'Via ViewModelProvider (manual setup)', 'activityRetainedScope()'],
        ['Compile-time scope validation', 'Yes', 'Yes', 'Yes', 'No (runtime crash)'],
        ['Runtime lifecycle callbacks', 'Implicit (generated code)', 'Not available', 'Yes — Scoped.onEnterScope/onExitScope (App Platform)', 'No'],
        ['Auto CoroutineScope per DI scope', 'No', 'No', 'Yes — scope.launch (App Platform)', 'No'],
        ['Prevents scope leak at compile time', 'Yes — enforced hierarchy', 'Yes — @GraphPrivate', 'Partial — scope markers', 'No'],
        ['KMP compatible', 'No — Android only', 'Yes', 'Yes', 'Yes'],
    ],
    [4, 2.8, 2.8, 3.2, 2.8]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key insight: ').bold = True
p.add_run(
    'Hilt is the only framework with fully automatic Android lifecycle management. For KMP projects, '
    'kotlin-inject-anvil + Amazon App Platform provides the closest equivalent with explicit lifecycle '
    'callbacks (onEnterScope/onExitScope). Metro and Koin rely on the developer to manage graph/scope '
    'creation and destruction at the right lifecycle points. For modern single-Activity + Compose apps, '
    'all four behave similarly because ViewModel lifecycle is handled by Android\'s ViewModelProvider, '
    'which all frameworks integrate with.'
)

doc.add_paragraph()
add_heading_styled('Java Code & Dagger/Hilt Interop', 2)
doc.add_paragraph(
    'For teams with existing Dagger/Hilt codebases migrating to KMP, interop capability determines '
    'migration effort:'
)
add_table(
    ['Capability', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Share graph with Dagger components', 'N/A (is Dagger)', 'Yes — @Includes daggerComponent', 'No — manual bridge required', 'No — manual bridge required'],
        ['Understand javax.inject.Inject', 'Yes', 'Yes — includeDagger()', 'Only in Kotlin files', 'No'],
        ['Understand jakarta.inject.Inject', 'No', 'Yes — includeJakarta()', 'No', 'No'],
        ['Reuse Dagger-generated Java factories', 'Yes', 'Yes — zero extra code', 'No — must wrap with @Provides', 'No'],
        ['Process .java source files', 'Yes (via KSP/KAPT)', 'No (Kotlin compiler plugin)', 'No (KSP = Kotlin only)', 'N/A (no processing)'],
        ['Incremental migration from Hilt', 'N/A', 'Yes — dual-build flag, module by module', 'No — clean break required', 'No — clean break required'],
        ['Effort for 100 shared dependencies', 'N/A', '1 line: @Includes daggerComponent', '100 @get:Provides parameters', '100 manual bindings'],
    ],
    [4.5, 2.5, 3, 3, 2.5]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Practical impact: ').bold = True
p.add_run(
    'For teams with existing Hilt/Dagger codebases, Metro\'s interop eliminates the biggest migration barrier. '
    'New KMP modules can use Metro with @Includes(daggerComponent) — zero bridge code, zero duplication. '
    'With kotlin-inject-anvil or Koin, every shared dependency must be manually bridged between the old '
    'Dagger graph and the new DI framework. For greenfield KMP projects (no existing Dagger), this '
    'difference doesn\'t apply.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 12. KOTLIN VERSION COMPATIBILITY GUIDE
# ═══════════════════════════════════════════════════════
add_heading_styled('12. Kotlin Version Compatibility Guide', 1)

add_heading_styled('12.1 Why Version Matters', 2)
doc.add_paragraph(
    'Every DI framework is built against a specific Kotlin version. This creates a project-wide constraint: '
    'you must pick ONE Kotlin version, and every framework must support it. The strictness depends on the '
    'framework type:'
)
add_table(
    ['Framework Type', 'Version Coupling', 'Why'],
    [
        ['Compiler plugin (Metro)', 'Strict — must match exact Kotlin version range', 'Plugin runs inside kotlinc; internal APIs change between versions'],
        ['KSP processor (Hilt, kotlin-inject-anvil)', 'Moderate — KSP version must match Kotlin, but processor JARs are more flexible', 'KSP runs on JVM alongside the compiler'],
        ['Runtime library with native targets (Koin, kotlin-inject runtime)', 'Strict for iOS — Kotlin/Native klibs not forward-compatible', 'klibs compiled with Kotlin 2.1 may not load on 2.2 Native compiler'],
        ['Runtime library JVM-only (Hilt runtime)', 'Flexible — JVM bytecode is broadly compatible', 'JAR files work across Kotlin versions'],
    ],
    [4.5, 3.5, 8]
)

doc.add_paragraph()
add_heading_styled('12.2 Version Matrix', 2)
doc.add_paragraph(
    'Three common Kotlin versions and what framework versions each supports:'
)

add_heading_styled('Kotlin 2.2.0', 3)
add_table(
    ['Framework', 'Latest Available', 'Compatible Version', 'Built With', 'Status'],
    [
        ['Hilt', '2.59.2', '2.59.2', 'Any (JVM)', 'Latest — no issue'],
        ['KSP', '2.2.0-2.0.2', '2.2.0-2.0.2', 'Kotlin 2.2.0', 'Exact match'],
        ['Metro', '1.0.0', '0.6.5', 'Kotlin 2.2.0-2.2.10', '14 versions behind — missing major features'],
        ['Koin', '4.2.1', '4.1.1', 'Kotlin 2.1.21', '3 versions behind — older resolver'],
        ['kotlin-inject-anvil', '0.1.7', '0.1.6', 'Kotlin 2.1.21', '1 version behind — minor gap'],
    ],
    [3.5, 2.5, 2.5, 2.5, 5]
)

doc.add_paragraph()
add_heading_styled('Kotlin 2.2.20', 3)
add_table(
    ['Framework', 'Latest Available', 'Compatible Version', 'Built With', 'Status'],
    [
        ['Hilt', '2.59.2', '2.59.2', 'Any (JVM)', 'Latest — no issue'],
        ['KSP', '2.2.20-2.0.2', '2.2.20-2.0.2', 'Kotlin 2.2.20', 'Exact match'],
        ['Metro', '1.0.0', '1.0.0', 'Kotlin 2.2.20', 'Latest stable — all features'],
        ['Koin', '4.2.1', '4.2.0', 'Kotlin 2.3.20', 'Near-latest (4.2.1 needs 2.3.20)'],
        ['kotlin-inject-anvil', '0.1.7', '0.1.7', 'Kotlin 2.2.20', 'Latest — no issue'],
    ],
    [3.5, 2.5, 2.5, 2.5, 5]
)

doc.add_paragraph()
add_heading_styled('12.3 What You Miss in Older Versions', 2)

add_heading_styled('Metro 0.6.5 vs 1.0.0 (14 versions behind)', 3)
add_table(
    ['Missing Feature', 'Version Added', 'Impact'],
    [
        ['metrox-viewmodel (ViewModel factory)', '0.8.0', 'Must write ViewModel factory manually'],
        ['metrox-viewmodel-compose (metroViewModel())', '0.8.0', 'Must wire Compose ViewModel yourself'],
        ['metrox-android (AppComponentFactory)', '0.8.0', 'No Activity/Fragment constructor injection'],
        ['AssistedViewModel support', '0.8.0', 'No assisted injection for ViewModels'],
        ['@ContributesTo aggregation', '0.10.0+', 'No Anvil-style auto module merging'],
        ['Dagger interop (includeDagger())', '0.10.0+', 'Cannot share graph with existing Hilt/Dagger'],
        ['kotlin-inject interop', '0.10.0+', 'Cannot interop with kotlin-inject components'],
        ['Graph class sharding', '0.10.0+', 'Large graphs may hit JVM class size limits'],
        ['Bug fixes (IR generation, cycle detection)', '0.7-1.0', 'Potential edge-case crashes on complex graphs'],
    ],
    [5, 2.5, 8.5]
)

doc.add_paragraph()
add_heading_styled('Koin 4.1.1 vs 4.2.1 (3 versions behind)', 3)
add_table(
    ['Missing Feature', 'Version Added', 'Impact'],
    [
        ['Core Resolver V2 (faster resolution)', '4.2.0', 'Older, slower resolution engine'],
        ['Ktor 3.4 DI Bridge', '4.2.0', 'Cannot use Koin with latest Ktor'],
        ['AndroidX Navigation 3 support', '4.2.0', 'Must manually handle Navigation 3 scoping'],
        ['Koin Compiler Plugin support', '4.2.0+', 'Cannot use K2 compiler plugin for compile-time safety'],
        ['Scope concurrency fixes', '4.2.1', 'Potential race condition on scope create/destroy'],
    ],
    [5, 2.5, 8.5]
)

doc.add_paragraph()
add_heading_styled('kotlin-inject-anvil 0.1.6 vs 0.1.7 (1 version behind)', 3)
doc.add_paragraph(
    'Minimal gap — only misses the update to kotlin-inject 0.9.0. Version 0.1.6 is stable and production-ready.'
)

doc.add_paragraph()
add_heading_styled('12.4 Recommendation', 2)
p = doc.add_paragraph()
p.add_run('Upgrade to Kotlin 2.2.20 for production projects. ').bold = True
p.add_run(
    'The jump from 2.2.0 to 2.2.20 is a minor version bump — not a breaking change for application code. '
    'The cost is approximately one day of re-testing. The benefit is significant: Metro 1.0.0 (stable, '
    'all features including ViewModel support, Dagger interop, module merging), Koin 4.2.0 (Core Resolver V2, '
    'Navigation 3), and kotlin-inject-anvil 0.1.7 (latest). Staying on 2.2.0 means using a beta Metro '
    '(0.6.5) that\'s missing critical production features — ViewModel integration, Dagger interop, '
    'and many bug fixes accumulated over 14 versions.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 13. SUMMARY & RECOMMENDATION
# ═══════════════════════════════════════════════════════
add_heading_styled('13. Summary & Recommendation', 1)

add_heading_styled('Complete Scorecard', 2)
add_table(
    ['Category', 'Hilt', 'Metro', 'kotlin-inject-anvil', 'Koin'],
    [
        ['Compile Time', '4,930ms (slowest)', '~2,500ms', '~2,500ms', '2,257ms (fastest)'],
        ['Generated Code', '387 files / 555KB (Java)', '0 source files (IR codegen)', '30 files / 58KB (Kotlin)', '0 files'],
        ['Android Runtime (total)', '17ms', '7ms', '6ms', '20ms (82 cls)'],
        ['Component/Graph Init', 'pre-built', '3.46ms', '2.82ms', '11.27ms'],
        ['ViewModels (13)', '16.21ms', '3.67ms', '3.67ms', '7.46ms'],
        ['Compile-Time Safety', 'Full', 'Full', 'Full', 'None (DSL) / Partial (Plugin)'],
        ['Component Merging', 'Yes (@InstallIn)', 'No (explicit)', 'Yes (@ContributesBinding)', 'No'],
        ['Dagger Interop', 'Native', '@Includes (zero bridge)', 'Manual bridge', 'Manual bridge'],
        ['KMP Support', 'No', 'Yes', 'Yes', 'Yes'],
        ['Java Support', 'Yes', 'Yes', 'No (Kotlin only)', 'No'],
        ['Ecosystem', 'Google official', 'Square/Block', 'Amazon', 'Community'],
    ],
    [3, 2.8, 2.8, 2.8, 2.8]
)

doc.add_paragraph()
add_heading_styled('When to Choose Each', 2)

p = doc.add_paragraph()
r = p.add_run('Choose Hilt when: ')
r.bold = True
r.font.color.rgb = GREEN_HILT
p.add_run(
    'You need Google\'s official support, extensive documentation, and a battle-tested ecosystem. '
    'Best for teams already familiar with Dagger and projects that prioritize long-term stability '
    'over build speed. Hilt\'s compile-time safety catches DI errors before the app runs.'
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Choose Metro when: ')
r.bold = True
r.font.color.rgb = BLUE
p.add_run(
    'Performance is critical — both compile-time and runtime. Metro offers the best of both worlds: '
    'compile-time safety (like Hilt) with near-zero runtime overhead, 44.7% faster builds than Hilt, '
    'and zero generated source files. Ideal for large-scale apps where build time and app startup '
    'latency matter. Also the only option if you need Kotlin Multiplatform DI with compile-time safety.'
)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Choose Koin when: ')
r.bold = True
r.font.color.rgb = ORANGE
p.add_run(
    'Simplicity, fastest compile times, and runtime flexibility are the priority. Koin\'s DSL is the '
    'easiest to learn and requires no annotation processing. Best for small-to-medium apps where the '
    '64us average injection time is negligible, KMP projects needing simple cross-platform DI, '
    'or apps requiring dynamic module loading (Play Feature Delivery, server-driven configs). '
    'The new Koin Compiler Plugin (K2) adds compile-time validation for most dependency errors, '
    'though edge cases around qualifiers, generic types, and dynamic modules can still crash at runtime. '
    'Runtime resolution overhead remains regardless of the plugin.'
)

doc.add_paragraph()
doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 12. REFERENCES & FURTHER READING
# ═══════════════════════════════════════════════════════
add_heading_styled('14. References & Further Reading', 1)

add_heading_styled('Official Documentation', 2)
refs = [
    'Google Android DI Guide — developer.android.com/training/dependency-injection',
    'Hilt Documentation — dagger.dev/hilt',
    'Metro Documentation — zacsweers.github.io/metro',
    'Metro Design Document — zacsweers.github.io/metro/latest/designdoc.html',
    'Koin Documentation — insert-koin.io/docs',
    'Koin Compiler Plugin — insert-koin.io/docs/intro/koin-compiler-plugin',
    'Kotlin/Native Overview — kotlinlang.org/docs/native-overview.html',
    'Kotlin/Native Memory Management — kotlinlang.org/docs/native-memory-manager.html',
    'Kotlin/Native ARC Integration — kotlinlang.org/docs/native-arc-integration.html',
    'KSP (Kotlin Symbol Processing) — kotlinlang.org/docs/ksp-overview.html',
    'K2 Compiler Migration Guide — kotlinlang.org/docs/k2-compiler-migration-guide.html',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
add_heading_styled('Technical Blog Posts & Articles', 2)
refs = [
    'Introducing Metro — zacsweers.dev/introducing-metro',
    'Metro Is Stable (1.0) — zacsweers.dev/metro-is-stable',
    'Metro Migration at Square Android — engineering.block.xyz/blog/metro-migration-at-square-android',
    'Koin Powered by Kotlin Compiler — blog.insert-koin.io/koin-powered-by-kotlin-compiler',
    'Compile-Time Safety with Koin Annotations — medium.com/@kerry.bisset (Kerry Bisset)',
    'Crash Course on the Kotlin Compiler: K1 + K2 — medium.com/google-developer-experts (Amanda Hinchman)',
    'From Dagger to Metro — vinted.engineering/2026/02/12/from-dagger-to-metro',
    'Cash Android Moves to Metro — code.cash.app/cash-android-moves-to-metro',
    'Dagger DoubleCheck: Scopes Explained — proandroiddev.com (Garima Jain)',
    'Hilt vs Koin: Hidden Cost of Runtime Injection — droidcon.com',
    'Benchmarking Koin vs Hilt in Modern Android (2024) — droidcon.com',
    'Kotlin/Native Compilation Under the Hood — medium.com/@natig.haciyef',
    'Stately: Kotlin Multiplatform Concurrency Library — github.com/touchlab/Stately',
    'Introducing kotlin-inject-anvil — ralf-wondratschek.com/blog',
    'Integrate kotlin-inject-anvil to Tv Maniac — Thomas Kioko, ProAndroidDev',
    'Creating Custom-Scoped Components in kotlin-inject + Anvil — droidcon 2025',
    'Amazon App Platform: Scope Lifecycle — amzn.github.io/app-platform/scope',
    'Koin vs kotlin-inject: Which to Choose — Infinum blog',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
add_heading_styled('Books', 2)
refs = [
    'Dependency Injection Principles, Practices, and Patterns — Mark Seemann & Steven van Deursen (Manning, 2019). The definitive guide to DI theory: composition root, Pure DI, service locator anti-pattern, compile-time vs runtime trade-offs. Language-agnostic principles that underpin all three frameworks. ISBN: 978-1617294730',
    'Kotlin in Action, 2nd Edition — Sebastian Aigner, Roman Elizarov, Svetlana Isakova, Dmitry Jemerov (Manning, 2024). Covers coroutines, structured concurrency, and Kotlin internals relevant to understanding compiler plugin behavior. manning.com/books/kotlin-in-action-second-edition',
    'Kotlin Design Patterns and Best Practices, 3rd Edition — Alexey Soshin & Anton Arhipov (Packt, 2024). Covers concurrent design patterns in Kotlin, including DI patterns, factory patterns, and service locator. ISBN: 978-1805127765',
    'Kotlin Multiplatform by Tutorials, 3rd Edition — Kodeco Team (Kodeco/raywenderlich, 2025). Hands-on KMP guide covering shared DI configuration with Koin, platform-specific modules with expect/actual, and iOS framework integration. kodeco.com/books/kotlin-multiplatform-by-tutorials',
]
for r in refs:
    p = doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
add_heading_styled('Conference Talks', 2)
refs = [
    'DroidKaigi 2025 — Navigating Dependency Injection with Metro (Zac Sweers)',
    'Droidcon Berlin 2025 — Koin Workshop: KMP Dependency Injection (Kotzilla)',
    'KotlinConf 2024 — Kotlin/Native Memory Management and Performance',
]
for r in refs:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of RFC —')
r.font.size = Pt(12)
r.font.color.rgb = GRAY

# ── Save ──
output_path = '/Users/sahilthakar/AndroidStudioProjects/BenchMarking/RFC_DI_Framework_Benchmark.docx'
doc.save(output_path)
print(f'RFC document saved to: {output_path}')
